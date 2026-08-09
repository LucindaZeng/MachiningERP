from __future__ import annotations

from pathlib import Path
from datetime import date
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from department_controls import DEPARTMENT_DETAILS, GLOBAL_ALERT_RULES, GLOBAL_APPROVAL_RULES
from order_lifecycle import (
    LIFECYCLE_REPORTS,
    LIFECYCLE_STAGES,
    ORDER_ALERTS,
    ORDER_TYPE_ALERTS,
    ORDER_TYPE_CONTROL_RULES,
    ORDER_TYPE_REPORTS,
    ORDER_TYPES,
    ORDER_MONITORING_END,
    ORDER_MONITORING_START,
    QUANTITY_HANDOFF_RULES,
    SLA_RULES,
    TIMING_FIELDS,
)
from rework_costing import (
    REWORK_ALERTS,
    REWORK_APPROVAL_RULES,
    REWORK_COST_COMPONENTS,
    REWORK_COST_PRINCIPLES,
    REWORK_ORDER_FIELDS,
    REWORK_REPORTS,
    REWORK_SCENARIOS,
)
from hong_kong_orders import (
    HK_ALERTS,
    HK_APPROVALS,
    HK_CALCULATION_RULES,
    HK_DEPARTMENT_CONTROLS,
    HK_DOCUMENTS,
    HK_ORDER_FIELDS,
    HK_ORDER_PRINCIPLES,
    HK_ORDER_WORKFLOW,
    HK_REPORTS,
)
from process_costing import (
    PROCESS_COST_ACCEPTANCE,
    PROCESS_COST_ALERTS,
    PROCESS_COST_APPROVALS,
    PROCESS_COST_COMPONENTS,
    PROCESS_COST_FIELDS,
    PROCESS_COST_FORMULAS,
    PROCESS_COST_MATRIX,
    PROCESS_COST_PRINCIPLES,
    PROCESS_COST_REPORTS,
    PROCESS_COST_WORKFLOW,
)


ROOT = Path(__file__).resolve().parents[1]
DOC_VERSION = "V2.1"
DOC_DATE = "2026年7月12日"
OUT = ROOT / "docs" / f"制造业ERP软件规划方案_{DOC_VERSION}.docx"
ASSET_DIR = ROOT / ".tmp" / "erp_plan_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D8DEE8"
WHITE = "FFFFFF"
BLACK = "111111"
GOLD = "B7791F"
RED = "9B1C1C"
GREEN = "276749"

TOC_TITLES = [
    "建设背景、目标与规划原则",
    "总体业务蓝图与系统架构",
    "订单全生命周期流程与节点时效监测",
    "组织、部门与岗位应用蓝图",
    "分阶段实施路线",
    "主数据、产品工程与PLM",
    "业务报价、合同、订单与客户协同",
    "采购、委外与供应商协同",
    "PMC、MRP与APS计划管理",
    "生产执行、设备与MES",
    "品质管理与QMS",
    "仓储、物流与WMS",
    "财务、业财一体化与制造成本",
    "系统单据与报表中心",
    "统一预警、待办与升级中心",
    "客户与供应商门户、安全和权限",
    "数据分析、指标与经营驾驶舱",
    "集成、数据、安全与非功能要求",
    "产品路线与厂商选型建议",
    "项目治理、实施、培训与验收",
    "端到端演示与验收场景",
    "附录A：功能需求编号矩阵",
    "附录B：系统生成文件清单",
    "附录C：选型演示问题清单",
    "附录D：厂商、行情及关务财税官方资料参考",
]

# Filled after the first pagination render; keeping this map in the source makes
# the static TOC deterministic in Word and in headless PDF conversion.
TOC_PAGES = {
    "建设背景、目标与规划原则": "4",
    "总体业务蓝图与系统架构": "6",
    "订单全生命周期流程与节点时效监测": "8",
    "组织、部门与岗位应用蓝图": "39",
    "分阶段实施路线": "89",
    "主数据、产品工程与PLM": "94",
    "业务报价、合同、订单与客户协同": "96",
    "采购、委外与供应商协同": "99",
    "PMC、MRP与APS计划管理": "103",
    "生产执行、设备与MES": "105",
    "品质管理与QMS": "108",
    "仓储、物流与WMS": "110",
    "财务、业财一体化与制造成本": "111",
    "系统单据与报表中心": "127",
    "统一预警、待办与升级中心": "142",
    "客户与供应商门户、安全和权限": "144",
    "数据分析、指标与经营驾驶舱": "145",
    "集成、数据、安全与非功能要求": "147",
    "产品路线与厂商选型建议": "148",
    "项目治理、实施、培训与验收": "150",
    "端到端演示与验收场景": "151",
    "附录A：功能需求编号矩阵": "159",
    "附录B：系统生成文件清单": "193",
    "附录C：选型演示问题清单": "195",
    "附录D：厂商、行情及关务财税官方资料参考": "197",
}

FONT_EAST_ASIA = "Arial Unicode MS"
FONT_ASCII = "Calibri"
_FONT_CANDIDATES = [
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
]
FONT_PATH = next((f for f in _FONT_CANDIDATES if Path(f).exists()), _FONT_CANDIDATES[0])


def set_run_font(run, size=None, bold=None, color=None, italic=None, east_asia=FONT_EAST_ASIA):
    run.font.name = FONT_ASCII
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_ASCII)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_ASCII)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C8D0DC", size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9, color=MUTED)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "在Word中按Ctrl+A后按F9可更新目录和页码。"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2, placeholder, fld_char3])


def add_update_fields_setting(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_picture_alt_text(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def make_font(size: int, bold=False):
    return ImageFont.truetype(FONT_PATH, size=size, index=0)


def draw_wrapped_center(draw, rect, text, font, fill, max_chars=12):
    lines = []
    current = ""
    for ch in text:
        current += ch
        if len(current) >= max_chars:
            lines.append(current)
            current = ""
    if current:
        lines.append(current)
    bbox = draw.multiline_textbbox((0, 0), "\n".join(lines), font=font, spacing=8, align="center")
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    x = rect[0] + (rect[2] - rect[0] - tw) / 2
    y = rect[1] + (rect[3] - rect[1] - th) / 2
    draw.multiline_text((x, y), "\n".join(lines), font=font, fill=fill, spacing=8, align="center")


def create_architecture_diagram(path: Path):
    img = Image.new("RGB", (1800, 1150), "white")
    draw = ImageDraw.Draw(img)
    title_font = make_font(48, True)
    label_font = make_font(32, True)
    body_font = make_font(28)
    draw.text((900, 50), "目标系统总体架构", font=title_font, fill="#203748", anchor="ma")
    bands = [
        ("外部协同层", "客户门户｜供应商SRM门户｜交易所/持牌金属行情｜在线报价｜交付计划｜ASN送货｜质量对账", "#EAF2F8"),
        ("经营管理层", "ERP：报价合同｜采购委外｜金属价格｜库存｜财务成本｜行政考勤/劳务｜外贸｜单据报表", "#E8EEF5"),
        ("计划与工程层", "PLM：图纸/BOM/工艺/变更    MRP：净需求与齐套    APS：有限产能排程与资源冲突", "#EEF5E9"),
        ("生产执行层", "MES：CNC/车床｜后工序｜独立组装｜一码到底｜不跳工序｜返工｜设备与人员绩效", "#FFF4E5"),
        ("质量与物流层", "QMS：IQC/IPQC/FQC/OQC/CAPA/追溯    WMS：库位/批次/余料/拣配/盘点/仓位图", "#FCECEC"),
        ("平台与数据层", "主数据｜预警中心｜工作流｜BI/AI分析｜指标语义层｜接口平台｜权限审计｜移动端｜备份与运维", "#F2F4F7"),
    ]
    y = 145
    for label, text, fill in bands:
        draw.rounded_rectangle((80, y, 1720, y + 135), radius=18, fill=fill, outline="#9AA8BA", width=3)
        draw.rounded_rectangle((95, y + 16, 345, y + 119), radius=14, fill="#203748")
        draw.text((220, y + 67), label, font=label_font, fill="white", anchor="mm")
        draw.text((390, y + 67), text, font=body_font, fill="#25364A", anchor="lm")
        if y < 820:
            draw.line((900, y + 135, 900, y + 158), fill="#697586", width=4)
            draw.polygon([(890, y + 153), (910, y + 153), (900, y + 165)], fill="#697586")
        y += 158
    img.save(path, quality=95)


def create_process_diagram(path: Path):
    img = Image.new("RGB", (2000, 880), "white")
    draw = ImageDraw.Draw(img)
    title_font = make_font(46, True)
    box_font = make_font(26, True)
    small_font = make_font(22)
    draw.text((1000, 45), "端到端业务主线与数据闭环", font=title_font, fill="#203748", anchor="ma")
    boxes = [
        ("客户需求与图纸", "门户/业务"),
        ("报价与成本", "CPQ/ERP"),
        ("合同评审与订单", "ERP/工作流"),
        ("产品工程与变更", "PLM"),
        ("MRP与齐套", "ERP/PMC"),
        ("采购/委外/排程", "SRM/APS"),
        ("扫码生产与质量", "MES/QMS"),
        ("入库交付与外贸", "WMS/ERP"),
        ("对账成本与分析", "财务/BI/AI"),
    ]
    x0, y0, bw, bh, gap = 60, 210, 185, 165, 31
    fills = ["#EAF2F8", "#E8EEF5", "#EEF5E9", "#FFF4E5", "#FCECEC"]
    for i, (title, sub) in enumerate(boxes):
        x = x0 + i * (bw + gap)
        fill = fills[i % len(fills)]
        draw.rounded_rectangle((x, y0, x + bw, y0 + bh), radius=18, fill=fill, outline="#7E8DA1", width=3)
        draw_wrapped_center(draw, (x + 8, y0 + 12, x + bw - 8, y0 + 105), title, box_font, "#203748", 7)
        draw.text((x + bw / 2, y0 + 138), sub, font=small_font, fill="#596579", anchor="mm")
        if i < len(boxes) - 1:
            ax = x + bw
            draw.line((ax + 4, y0 + bh / 2, ax + gap - 7, y0 + bh / 2), fill="#697586", width=4)
            draw.polygon([(ax + gap - 14, y0 + bh / 2 - 8), (ax + gap - 14, y0 + bh / 2 + 8), (ax + gap - 3, y0 + bh / 2)], fill="#697586")
    draw.rounded_rectangle((180, 515, 1820, 750), radius=20, fill="#F4F6F9", outline="#B4BECA", width=3)
    draw.text((1000, 555), "贯穿全流程的控制能力", font=box_font, fill="#203748", anchor="ma")
    control_text = (
        "十三部门组织权限｜统一编码与版本｜一码到底｜过程预警｜系统生成单据/报表｜客户与供应商隔离｜"
        "质量冻结与返工成本｜考勤与劳务结算｜财务和老板AI｜订单、批次、工序、人员与凭证全链追溯"
    )
    draw_wrapped_center(draw, (220, 585, 1780, 730), control_text, small_font, "#334155", 42)
    img.save(path, quality=95)


def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_ASCII
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_ASCII)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_ASCII)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT_ASCII
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_ASCII)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_ASCII)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    styles["Heading 1"].paragraph_format.page_break_before = False

    for list_name in ("List Bullet", "List Number"):
        st = styles[list_name]
        st.font.name = FONT_ASCII
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_ASCII)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_ASCII)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.375)
        st.paragraph_format.first_line_indent = Inches(-0.194)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.208

    if "Caption Custom" not in styles:
        cap = styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Caption Custom"]
    cap.font.name = FONT_ASCII
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    cap.font.size = Pt(9.5)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)

    # Running header and footer, intentionally quiet for a formal planning document.
    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("ERP软件规划方案  |  订单驱动型机加工与零部件制造")
    set_run_font(r, size=9, color=MUTED)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D8DEE8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = fp.add_run("内部规划文件  |  第 ")
    set_run_font(rr, size=9, color=MUTED)
    add_page_field(fp)
    rr2 = fp.add_run(" 页")
    set_run_font(rr2, size=9, color=MUTED)

    doc.core_properties.title = "制造业ERP软件规划方案"
    doc.core_properties.subject = "订单驱动型机加工与零部件制造企业数字化管理系统规划"
    doc.core_properties.author = "ERP规划项目组"
    doc.core_properties.keywords = "ERP, PLM, APS, MES, QMS, WMS, SRM, CNC, 后工序, 组装, FAI, 行政考勤, 劳务结算, 金属材料价格, AI分析"
    add_update_fields_setting(doc)
    return doc


def add_para(doc, text: str, bold_prefix: str | None = None, italic=False, align=None, after=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    return p


def add_bullets(doc, items: Iterable[str], level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        r = p.add_run(item)
        set_run_font(r)


def _new_decimal_num_id(doc) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered(doc, items: Iterable[str], font_size=None, space_after=4, line_spacing=1.208):
    num_id = _new_decimal_num_id(doc)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_el])
        p_pr.append(num_pr)
        r = p.add_run(item)
        set_run_font(r, size=font_size)


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    return p


def add_callout(doc, title: str, body: str, fill=LIGHT_GRAY, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r1 = p.add_run(f"{title}  ")
    set_run_font(r1, bold=True, color=color)
    r2 = p.add_run(body)
    set_run_font(r2, color=BLACK)


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_dxa: Sequence[int], font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h))
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if i == 0 and len(headers) <= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run_font(r, size=font_size)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path: Path, caption: str, alt: str, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_picture_alt_text(shape, caption, alt)
    cp = doc.add_paragraph(caption, style="Caption Custom")
    cp.paragraph_format.keep_with_next = False


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("数字化管理与智能制造规划")
    set_run_font(r, size=12, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("制造业ERP软件规划方案")
    set_run_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("适用于订单驱动型机加工与零部件制造企业")
    set_run_font(r, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(65)
    r = p.add_run("覆盖 ERP / PLM / MRP / APS / MES / QMS / WMS / SRM / 行政考勤 / 客户门户 / BI / AI分析问答")
    set_run_font(r, size=10.5, italic=True, color=MUTED)

    metadata = [
        ("文档版本", DOC_VERSION),
        ("编制日期", DOC_DATE),
        ("文档性质", "ERP选型、实施规划与验收依据"),
        ("适用范围", "报价、订单、工程、计划、采购、金属行情、委外、生产、后工序、品质、仓储、财务、行政考勤、外贸、AI分析及协同"),
    ]
    add_table(doc, ["项目", "内容"], metadata, [2100, 7260], font_size=10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("内部规划文件")
    set_run_font(r, size=10, bold=True, color=MUTED)
def add_front_matter(doc):
    add_heading(doc, "文档说明", 1)
    add_para(doc, "本方案依据订单驱动型机加工与零部件制造企业的具体管理需求编制。内容覆盖前期沟通中提出的全部业务功能，并进一步展开到业务规则、数据字段、系统边界、自动单据、门户协同、预警、权限、实施与验收。")
    add_callout(doc, "V2.0增补", "保留V1.9全部工序级成本能力，并按最新确认简化香港代生产规则：只在客户主数据勾选特定客户；只有该客户的正式业务订单按输入价格70%自动计价，样品、模具及其他客户保持原价，价格以外全部流程不变。", fill="EEF5E9", color=GREEN)
    add_callout(doc, "V2.1增补", "查漏补缺增补：订单交付后尾数返工与入库四条处理路径及结案数量平衡校验；长时间工单未完结账龄监控与强制完结评审；PMC架机表自动生成、调整留痕、架机率与平衡化生产；中途改工序控制；财务与老板AI问答明确采用DeepSeek V4 Pro模型。", fill="EAF2FB", color=DARK_BLUE)
    add_callout(doc, "使用说明", "本方案可直接作为ERP项目立项材料、需求规格初稿、供应商RFP附件、产品演示脚本及实施验收基线。最终合同需将关键功能、接口、报表和验收数据逐项固化。")
    add_table(
        doc,
        ["使用对象", "主要用途", "重点关注"],
        [
            ("总经办/管理层", "确认建设范围、优先级和全厂经营", "订单、交付、产能、成本、库存、质量、现金和人员指标"),
            ("业务部门", "确认报价、合同、订单和客户协同", "报价效率、版本、交期承诺和客户可见进度"),
            ("工程/PMC", "确认图纸、BOM、工艺、MRP和排程", "齐套、产能、资源冲突、合批和变更影响"),
            ("采购/供应链", "确认供应商询报价、交付和对账", "在线报价、订单确认、ASN、送货单、质量与绩效"),
            ("生产/后工序/品质/仓库", "确认扫码、返工、组装、全检、包装和库位", "CNC/车床、后工序、组装工单、质量冻结、余料与仓位图"),
            ("财务/IT", "确认业财一体化、成本、AI问答、接口和安全", "自动凭证、制造成本、月结、资金、AI分析可信度和数据隔离"),
            ("行政", "确认人员、班次、考勤、异常和劳务结算", "考勤采集、日报月报、多劳务规则、对账与隐私权限"),
        ],
        [1500, 3300, 4560],
        font_size=9.2,
    )

    add_heading(doc, "目录", 1)
    half = (len(TOC_TITLES) + 1) // 2
    left = TOC_TITLES[:half]
    right = TOC_TITLES[half:]
    toc_rows = []
    for idx in range(half):
        left_title = left[idx]
        right_title = right[idx] if idx < len(right) else ""
        toc_rows.append((left_title, TOC_PAGES.get(left_title, ""), right_title, TOC_PAGES.get(right_title, "") if right_title else ""))
    add_table(doc, ["章节", "页码", "章节", "页码"], toc_rows, [3840, 600, 4320, 600], font_size=8.6)
    add_para(doc, "说明：目录页码依据本版文档排版生成；在Word中修改正文后，应同步更新目录页码。", italic=True, after=2)
    doc.add_page_break()


def section_overview(doc):
    add_heading(doc, "建设背景、目标与规划原则", 1)
    add_heading(doc, "业务特征判断", 2)
    add_para(doc, "从需求可以判断，企业具有按客户图纸报价、多品种小批量、工序多、机台和夹治具约束明显、部分工序委外、交付压力高、质量追溯严格等特征。管理对象不是单纯的商品库存，而是订单、图纸版本、材料批次、工艺路线、工序在制、设备、人员、委外供应商、质量异常和制造成本的动态组合。")
    add_bullets(doc, [
        "订单来源包含样品、批量和备料，报价数量变化会影响材料利用率、调机、工装摊销和委外最低批量费用。",
        "工程变更频繁，图纸、BOM、工艺路线、检验标准和在制订单必须保持版本一致。",
        "PMC需要滚动掌握现有库存、已分配量、在制、在途、冻结、不良、损耗和安全库存，实时计算齐套和净需求。",
        "金属材料价格波动直接影响采购时点、报价有效期、标准成本、库存价值、订单毛利和现金需求，需要统一行情与价格快照。",
        "生产现场要求一码到底、扫码开完工、不跳工序、条件跳序审批、跨工序退回、内部返工与复检。",
        "供应商与客户需要通过网页参与询报价、订单确认、进度、交货、质量、文件和对账协同。",
        "公司内部由财务、总经办、业务、工程、PMC、采购、委外、仓库、生产、品质、后工序、行政和IT十三部门协同，既需要全厂统一数据，又必须明确部门数据责任和权限。",
        "行政需要汇总多来源考勤、班次、请假加班和劳务用工，并按不同劳务合同规则生成报告、对账和结算。",
        "所有正式文件和运营报表必须由系统数据自动生成，不允许再以人工维护Excel作为事实来源。",
    ])

    add_heading(doc, "总体建设目标", 2)
    add_numbered(doc, [
        "建立从客户图纸、报价、合同、订单、工程、计划、供应链、生产、质量、仓储、交付、财务到经营分析的端到端数据闭环。",
        "建立唯一的物料、产品、图纸、BOM、工艺、客户、供应商、设备、工装、检具和仓位主数据。",
        "让PMC基于真实库存、订单和进度自动计算MRP、齐套、采购、委外和生产建议。",
        "让车间执行数据通过扫码或设备采集产生，使工序进度、良率、返工、报废和停机分析可追溯。",
        "让供应商和客户在授权门户上查看自己的订单、项目、交付计划和文件，并在线完成报价、确认、送货、质量及对账操作。",
        "让报价单、成本分析表、外贸文件、进度表、对账单和分析报表由系统按模板、流程和版本自动生成。",
        "通过预警中心对报价、审批、采购、委外、工序、检验、返工、设备和交付超期进行分级提醒与升级。",
        "形成报价成本、标准成本和实际成本的对比，识别材料、人工、机时、委外、返工、报废和质量损失。",
        "建立采购、财务和业务共享但按字段隔离的金属材料价格中心，使每次采购、报价和成本分析使用可追溯的价格与汇率快照。",
        "为十三部门建立独立工作台、责任、审批、报表和预警，使CNC/车床、后工序、独立组装、质量与行政劳务均进入统一数据链。",
    ])

    add_heading(doc, "规划原则", 2)
    add_table(
        doc,
        ["原则", "具体要求"],
        [
            ("业务闭环优先", "先确保订单、物料、生产、质量、交付和财务能够闭环，再建设高级算法和AI功能。"),
            ("数据一次产生", "业务发生时直接在责任岗位录入或扫码采集，后续单据、报表和凭证引用同一来源。"),
            ("过程与结果并重", "ERP管理资源与核算，MES/WMS/QMS控制现场过程，不能只做月末结果统计。"),
            ("标准产品优先", "优先采用成熟模块和可配置流程；相似产品报价、图纸识别、套料及合批优化列入专项定制。"),
            ("版本和追溯优先", "图纸、BOM、工艺、报价、订单、合同和正式文件全部保留版本与生效范围。"),
            ("外部可见受控", "客户看到经发布的里程碑，供应商只看到自己的询价、订单和质量数据，内部成本与其他合作方信息严格隔离。"),
            ("分阶段验收", "每阶段以真实订单跑通并达到数据质量门槛后再进入下一阶段。"),
        ],
        [1800, 7560],
        font_size=9.3,
    )


def section_architecture(doc):
    add_heading(doc, "总体业务蓝图与系统架构", 1)
    arch = ASSET_DIR / "architecture.png"
    process = ASSET_DIR / "process.png"
    create_architecture_diagram(arch)
    create_process_diagram(process)
    add_figure(doc, arch, "图1 目标系统总体架构", "从外部协同、经营管理、计划工程、生产执行、质量物流到平台数据的六层系统架构。")
    add_heading(doc, "系统边界", 2)
    add_table(
        doc,
        ["系统", "主要职责", "本项目关键范围"],
        [
            ("ERP", "经营交易、资源计划、库存、财务和成本", "报价、合同、销售、采购、委外、MRP、库存、外贸、应收应付、制造成本和自动凭证"),
            ("PLM", "产品与工程技术数据", "客户图档、物料、BOM、工艺、作业指导书、工装刀具、图纸版本和工程变更"),
            ("APS", "有限产能排程和交期模拟", "机台、人员、夹具、材料、委外、换型、优先级、合批及资源冲突"),
            ("MES", "车间派工、执行和数据采集", "一码到底、扫码报工、不跳工序、退序返工、异常、绩效、设备状态和OEE"),
            ("QMS", "检验、不合格、改进与追溯", "IQC、首件、巡检、完工、出货、NCR、MRB、CAPA、8D、检具和质量成本"),
            ("WMS", "仓库现场过程控制", "ASN收货、上架、库位、批次、余料、备料、拣配、盘点、仓位图和条码防错"),
            ("SRM门户", "供应商外部协同", "在线询报价、订单确认、交期、生产/委外进度、ASN、送货单、质量、对账、发票和付款查询"),
            ("金属价格中心", "外部行情与企业材料价格", "交易所/持牌数据、供应商报价、历史成交、落地价换算、价格快照、角色视图和预警"),
            ("客户门户", "客户项目与交付协同", "报价/订单文件、受控里程碑、交货计划、发货、质量文件、客诉、返工和收货确认"),
            ("行政/考勤", "人员、班次、考勤和劳务结算", "组织岗位、打卡、请假加班、异常、考勤报表、多劳务费率、对账和财务应付"),
            ("BI/报表中心", "统一指标和系统生成文件", "进度、预警、成本、质量、设备、库存、供应商绩效以及PDF/Excel版本管理"),
            ("AI分析问答", "财务与老板自然语言分析", "受权限控制的指标问答、原因下钻、经营模拟、引用单据、分析报告和完整审计"),
        ],
        [1320, 2760, 5280],
        font_size=8.7,
    )
    add_figure(doc, process, "图2 端到端业务主线与数据闭环", "从客户需求、报价、订单和工程到计划、采购、生产、质量、交付、结算与分析的端到端闭环。")

    add_heading(doc, "核心业务对象及关联", 2)
    add_bullets(doc, [
        "客户需求、询价、客户图纸和样品要求关联到报价版本与成本分析。",
        "报价审批后转销售订单，销售订单关联合同评审、交货批次、产品版本和客户可见计划。",
        "产品关联原材料规格、BOM、工艺路线、设备能力、夹治具、刀具、检验标准和包装标准。",
        "销售订单经MRP生成计划订单、采购申请、委外建议和生产工单；APS把任务分配到设备、班次和人员。",
        "工单通过批次码或序列码贯穿领料、工序、检验、返工、入库和发货。",
        "采购与委外订单通过供应商门户完成报价、确认、进度、ASN、送货、质量和对账。",
        "所有交易自动进入存货、应收、应付和制造成本，形成报价成本、标准成本和实际成本对比。",
    ])


def section_order_lifecycle(doc):
    add_heading(doc, "订单全生命周期流程与节点时效监测", 1)
    add_callout(
        doc,
        "核心流程基线",
        "新客户报价→客户确认→客户主数据→BOM申请→工程建BOM/内部图纸转换→程序并行准备→业务订单→业务经理审核→财务审核→订单评审→PMC运行MRP→采购下单→财务审核采购单→原材料到厂→仓库收料→IQC→可选切料→CNC/车床→逐工序品质门禁与数量交接→打磨去毛刺→品质→仓库外发委外表处→委外回料→品质→镭雕→品质→丝印→品质→可选独立组装→100%全检→包装。",
        fill="EEF5E9",
        color=GREEN,
    )
    add_heading(doc, "监测起止与两层时钟", 2)
    add_para(doc, f"主流程开始：{ORDER_MONITORING_START}", bold_prefix="主流程开始：")
    add_para(doc, f"生产履约结束：{ORDER_MONITORING_END}", bold_prefix="生产履约结束：")
    add_bullets(doc, [
        "报价、客户主数据和工程准备虽然发生在T0前，也必须记录自己的任务时效，用于报价响应和工程准备分析；但它们不混入已下业务订单后的主流程SLA。",
        "订单级总时钟回答从下业务订单到全检包装用了多久；节点级时钟回答某一次审核、MRP、采购、回料、检验、生产或委外用了多久。",
        "所有节点都要区分排队等待、首次响应、实际处理、经批准暂停、退回补充和总历时，避免把等待时间错误算成加工时间，或用暂停掩盖超期。",
    ])

    add_heading(doc, "业务订单三分类与收费规则", 2)
    add_callout(doc, "强制规则", "模具订单和样品订单不代表自动免费：二者可以收费，也可以按审批免费或部分收费；正式业务订单必须收费，即使采用账期、分期或月结，也必须形成合同应收义务。", fill="FFF4E5", color=GOLD)
    add_table(doc, ["订单类型", "客户付费规则", "可选收费方式", "审核与业务控制", "成本及财务处理"], ORDER_TYPES, [1100, 1250, 1900, 2700, 2410], font_size=7.4)
    add_heading(doc, "订单分类通用控制", 3)
    add_numbered(doc, ORDER_TYPE_CONTROL_RULES, font_size=9.2, space_after=3, line_spacing=1.12)
    add_heading(doc, "订单分类专项预警", 3)
    add_bullets(doc, ORDER_TYPE_ALERTS)
    add_heading(doc, "订单分类专项报表", 3)
    add_table(doc, ["报表", "系统生成内容"], ORDER_TYPE_REPORTS, [2500, 6860], font_size=8.6)

    add_heading(doc, "香港代生产客户正式订单70%价格规则", 2)
    add_callout(doc, "唯一触发条件", "仅当订单所选客户已在客户主数据中勾选“香港代生产价格客户”，并且订单类型为“正式业务订单”时，系统才将业务输入价格自动乘70%。样品订单、模具订单、未勾选客户均保持原输入价格；除价格换算外所有流程不变。", fill="FFF4E5", color=GOLD)
    add_callout(doc, "流程边界", "该规则不是新的订单模式，也不产生独立分成、特殊结算或特殊报关流程。系统只保存原始输入价格、70%系数和计算后正式订单价格，后续继续走普通正式订单的审核、MRP、采购、生产、品质、委外、报关、开票和收款流程。", fill="EEF5E9", color=GREEN)
    add_heading(doc, "客户价格规则核心原则", 3)
    add_numbered(doc, HK_ORDER_PRINCIPLES, font_size=9.1, space_after=3, line_spacing=1.12)
    add_heading(doc, "客户勾选与订单价格必填字段", 3)
    add_table(doc, ["数据分组", "字段与统一口径"], HK_ORDER_FIELDS, [2050, 7310], font_size=8.35)
    add_heading(doc, "70%价格计算与适用性校验", 3)
    add_table(doc, ["指标", "计算规则", "系统控制"], HK_CALCULATION_RULES, [1800, 3650, 3910], font_size=8.0)
    add_heading(doc, "香港代生产完整流程", 3)
    add_numbered(doc, HK_ORDER_WORKFLOW, font_size=9.0, space_after=3, line_spacing=1.12)
    add_heading(doc, "系统生成文件与资料包", 3)
    add_table(doc, ["文件/报表", "系统生成内容"], HK_DOCUMENTS, [2500, 6860], font_size=8.35)
    add_heading(doc, "部门责任与权限边界", 3)
    add_table(doc, ["角色", "责任与边界"], HK_DEPARTMENT_CONTROLS, [2150, 7210], font_size=8.35)
    add_heading(doc, "审核与授权权限", 3)
    add_numbered(doc, HK_APPROVALS, font_size=9.0, space_after=3, line_spacing=1.12)
    add_heading(doc, "客户价格规则专项预警", 3)
    add_bullets(doc, HK_ALERTS)
    add_heading(doc, "客户价格规则专项报表", 3)
    add_table(doc, ["报表", "系统生成内容"], HK_REPORTS, [2500, 6860], font_size=8.35)

    add_heading(doc, "节点计时数据模型", 2)
    add_table(doc, ["字段", "统一口径"], TIMING_FIELDS, [2200, 7160], font_size=8.8)
    add_heading(doc, "SLA、改期与升级规则", 2)
    add_numbered(doc, SLA_RULES, font_size=9.2, space_after=3, line_spacing=1.12)

    phase = None
    for item in LIFECYCLE_STAGES:
        if item["phase"] != phase:
            phase = item["phase"]
            add_heading(doc, phase, 2)
        add_heading(doc, f"{item['code']}  {item['stage']}", 3)
        add_para(doc, f"责任岗位：{item['owner']}；适用性：{item['required']}。", bold_prefix="责任岗位：")
        add_para(doc, f"计时起点：{item['start']}。", bold_prefix="计时起点：")
        add_para(doc, f"阶段动作：{item['actions']}", bold_prefix="阶段动作：")
        add_para(doc, f"完成口径：{item['finish']}。", bold_prefix="完成口径：")
        add_para(doc, f"系统控制：{item['control']}", bold_prefix="系统控制：")
        add_para(doc, f"预警重点：{item['alerts']}", bold_prefix="预警重点：")

    add_heading(doc, "工序数量交接与品质职责分离", 2)
    add_callout(doc, "示例", "CNC申报转移100件、打磨实际点收94件时，系统先转移94件；差异6件仍显示在CNC工序，立即向CNC责任人预警并进入数量差异闭环。品质只确认所检对象是否合格，不确认交接数量。", fill="FFF4E5", color=GOLD)
    add_numbered(doc, QUANTITY_HANDOFF_RULES, font_size=9.2, space_after=3, line_spacing=1.12)

    add_heading(doc, "返工子订单拆分、重复成本与责任归集", 2)
    add_callout(doc, "核心判断", "返工“由谁执行”和“由谁承担责任”是两个维度。返工重新经过的每道工序都要产生第二次成本；供应商是否需要新增付款取决于供应商是否为责任方，但即使供应商无偿返工，也要核算其给本订单造成的全部企业损失。", fill="FCECEC", color=RED)
    add_heading(doc, "返工订单必填数据", 3)
    add_table(doc, ["数据分组", "字段与统一口径"], REWORK_ORDER_FIELDS, [2100, 7260], font_size=8.5)
    add_heading(doc, "返工成本组成", 3)
    add_table(doc, ["成本组", "必须归集的内容"], REWORK_COST_COMPONENTS, [2100, 7260], font_size=8.5)
    add_heading(doc, "返工成本通用原则", 3)
    add_numbered(doc, REWORK_COST_PRINCIPLES, font_size=9.1, space_after=3, line_spacing=1.12)
    add_heading(doc, "两类关键责任场景", 3)
    for scenario, route, costs, payable, attribution in REWORK_SCENARIOS:
        add_heading(doc, scenario, 3)
        add_para(doc, f"受控返工路线：{route}", bold_prefix="受控返工路线：")
        add_para(doc, f"重复成本：{costs}", bold_prefix="重复成本：")
        add_para(doc, f"供应商追加应付：{payable}", bold_prefix="供应商追加应付：")
        add_para(doc, f"责任成本归集：{attribution}", bold_prefix="责任成本归集：")
    add_heading(doc, "返工审批、结算与关闭", 3)
    add_numbered(doc, REWORK_APPROVAL_RULES, font_size=9.1, space_after=3, line_spacing=1.12)
    add_heading(doc, "返工成本专项预警", 3)
    add_bullets(doc, REWORK_ALERTS)
    add_heading(doc, "返工成本专项报表", 3)
    add_table(doc, ["报表", "系统生成内容"], REWORK_REPORTS, [2450, 6910], font_size=8.5)

    add_heading(doc, "尾数返工与入库（订单交付后收尾）", 2)
    add_para(doc, "业务订单主批次交付后，剩余尾数（待返工不良品、返工合格品、超交余量和确认报废品）必须进入系统化收尾流程，禁止账外滞留。主批次发货过账后系统自动盘点该订单在制、待检、冻结和已完工未发货数量，生成订单尾数清单。")
    add_heading(doc, "尾数处理路径", 3)
    add_numbered(doc, [
        "返工后补交客户：生成返工子订单，复用返工拆单与责任成本规则，返工合格后按补交发货流程出货并计入原订单交付数量。",
        "返工后入库：客户需求已满足时，返工合格品按半成品或成品入库转为公司库存，入库成本取该批实际工序累计成本，后续订单优先占用。",
        "直接入库：无质量问题的超交余量直接入成品库，并与产品、图纸版本和材料批次关联，供后续订单或补货使用。",
        "报废：无法修复的尾数走报废审批，计入订单报废成本和责任分析。",
    ])
    add_heading(doc, "尾数控制规则", 3)
    add_bullets(doc, [
        "订单结案前系统强制校验：投入数量＝交付数量＋尾数入库数量＋报废数量＋在制返工数量，不平不得结案。",
        "尾数入库保留原订单号、批次码和完整工序履历，一码到底不中断；后续订单占用尾数库存时自动带出谱系。",
        "尾数返工成本按返工子订单独立归集，不冲抵主订单已确认成本；入库尾数按实际累计成本估价并与财务存货账一致。",
        "图纸或工艺版本升级后，旧版本尾数库存自动标记版本受限，占用前需工程确认可用性。",
    ])
    add_heading(doc, "尾数预警与报表", 3)
    add_bullets(doc, [
        "预警：订单发货后超期（默认7天）尾数未定路径提醒PMC和业务；尾数返工子订单超期按统一工序规则升级；尾数库存超保留期（默认90天）未被占用转入呆滞物料分析。",
        "报表：订单尾数清单（订单、产品、数量、状态、路径、责任人、截止时间）、尾数入库与占用履历表、尾数返工成本表（并入返工子订单成本明细口径）。",
    ])

    add_heading(doc, "长时间工单未完结监控", 2)
    add_para(doc, "除单工序超期预警外，系统从工单整体维度监控所有已下达未关闭的生产工单、返工子订单、委外单、采购单和尾数处理单，防止在制、成本和交期黑洞。")
    add_bullets(doc, [
        "账龄＝当前时间－工单下达时间；停滞时长＝当前时间－最后一次有效业务动作（报工、转序、检验、收发料）时间；账龄超过标准制造周期×系数（默认1.5）即视为超龄。",
        "账龄分层（默认可配置）：停滞超3天或账龄超标准周期1.2倍为关注（黄色）；停滞超7天或1.5倍为超龄（橙色，必须填原因和预计完结时间）；停滞超15天或2倍为严重超龄（红色，升级部门负责人并进入总经办超期看板）；停滞超30天为疑似死单（深红色，强制完结评审）。",
        "疑似死单由PMC发起完结评审，生产、品质、财务会签，判定剩余数量走返工、入库、报废或取消；关闭前校验数量平衡和工序成本完整性，不允许直接删除工单。",
        "报表：未完结工单账龄表（按部门、机台、订单类型、账龄分层，可下钻工单和最后动作）、停滞原因分析、月结前应关未关清单（作为工序成本月结检查前置输入）。",
    ])

    add_heading(doc, "订单级实时预警", 2)
    add_bullets(doc, ORDER_ALERTS)
    add_heading(doc, "订单生命周期系统报表", 2)
    add_table(doc, ["报表/看板", "系统生成内容"], LIFECYCLE_REPORTS, [2400, 6960], font_size=8.6)
    add_callout(doc, "工序成本同步要求", "订单进入切料、CNC、车床、打磨、清洗、振磨、委外、镭雕、丝印、组装、检验、全检或包装任一工序时，系统同时建立独立工序成本卡；报工完成、质量判定和数量交接后更新本工序新增成本及累计成本。", fill="FFF4E5", color=GOLD)
    add_callout(
        doc,
        "量化要求",
        "所有审核和业务节点必须至少输出平均值、中位数、P90、P95、按时完成率、超期率、首次响应时长、排队时长、实际处理时长和积压数量；能够按订单、客户、产品、部门、岗位、审核人、供应商、设备、工序、日期和紧急等级下钻。",
        fill="E8EEF5",
        color=DARK_BLUE,
    )


def add_department_control_detail(doc, department_name):
    detail = next(item for item in DEPARTMENT_DETAILS if item["name"] == department_name)
    add_callout(doc, "部门控制边界", detail["scope"], fill="F4F6F9", color=DARK_BLUE)
    add_heading(doc, "具体功能明细", 3)
    add_bullets(doc, detail["functions"])
    add_heading(doc, "审核与授权权限", 3)
    add_bullets(doc, detail["approvals"])
    add_heading(doc, "预警信息", 3)
    add_bullets(doc, detail["alerts"])
    add_heading(doc, "系统生成报表", 3)
    add_bullets(doc, detail["reports"])


def section_department_blueprint(doc):
    add_heading(doc, "组织、部门与岗位应用蓝图", 1)
    add_callout(doc, "组织原则", "ERP中的“部门”用于权限、责任、预算和核算，“工作中心”用于产能、派工、报工与成本。生产部、后工序部和组装不能只建成一个笼统生产部门；采购与委外也必须区分岗位和审批责任。")

    add_heading(doc, "组织、岗位、工作中心与数据权限", 2)
    add_bullets(doc, [
        "组织层级建议至少包含公司/法人、工厂、部门、班组、工作中心、成本中心和利润中心；同一对象可按规则建立映射，但不能用一个字段混代所有层级。",
        "部门范围为财务部、总经办、生产部、品质部、后工序部、工程部、PMC、采购、委外和行政；仓库、业务/销售、设备等岗位若暂不独立成部门，仍需作为明确职能和权限角色存在。",
        "生产部下设CNC和车床工作中心，并可继续按设备组、班次或区域细分；后工序部按清洗、丝印、镭雕、打磨、振磨、全检包装细分，组装单独建立工作中心。",
        "员工可有主岗位、兼任岗位、所属组织、班组、技能、设备授权、审批代理和有效期；兼任权限到期后自动收回。",
        "角色权限同时控制菜单、操作、组织范围、数据行、敏感字段、单据状态和审批额度；总经办可以分析全厂，但不等于可以修改财务凭证、检验结果或工序履历。",
        "每个部门建立首页工作台，展示本部门待办、今日任务、超期、异常、关键指标、需协同事项和可生成文件；数据来自源业务，不允许部门另建平行Excel台账。",
    ])
    add_heading(doc, "全公司统一审核权限规则", 2)
    add_bullets(doc, GLOBAL_APPROVAL_RULES)
    add_heading(doc, "全公司统一预警处理规则", 2)
    add_bullets(doc, GLOBAL_ALERT_RULES)
    add_callout(doc, "清单阅读方式", "下面每个部门都按四个维度展开：具体功能明细、审核与授权权限、预警信息、系统生成报表。审批阈值不在规划阶段写死具体金额，由项目蓝图阶段根据公司授权制度配置；但申请、复核、批准、执行和记账/放行的职责边界必须在上线前确定。", fill="EEF5E9", color=GREEN)
    add_table(
        doc,
        ["部门", "主要系统工作台", "数据与流程责任", "关键输出"],
        [
            ("财务部", "总账/往来/资金/成本/税务/AI", "财务主数据、凭证、应收应付、资金、资产、成本、月结", "全盘账、往来账、成本账、报表、资金与经营分析"),
            ("总经办", "经营驾驶舱/审批/预警/AI", "全厂指标口径、目标、重大审批、经营风险和改善闭环", "经营日报/月报、交付/现金/质量/产能分析和管理指令"),
            ("生产部", "MES/设备/班组", "CNC与车床派工、扫码、报工、设备、人员、异常和在制", "机台任务、产量、工时、良率、达成率、OEE和停机分析"),
            ("品质部", "QMS/检验/计量", "首检、点检/巡检、末件检、FAI、异常、放行和追溯", "电子检验记录、FAI报告、异常/CAPA和质量分析"),
            ("后工序部", "MES后工序/组装/包装", "清洗、丝印、镭雕、打磨、振磨、全检包装及独立组装", "后工序履历、全检记录、包装记录、组装工单与测试记录"),
            ("工程部", "PLM/工程/打样", "物料工程属性、BOM、图纸、程序、工艺、工装、变更和样品", "受控BOM/图纸/程序、工艺路线、样品和工程变更"),
            ("PMC", "MRP/APS/齐套/进度", "需求计划、物料齐套、生产/委外计划、优先级和交付协调", "采购/生产/委外建议、缺料、齐套、排程和交付进度"),
            ("采购", "SRM/采购/金属价格", "供应商、询比价、定标、采购订单、到货和采购进度", "询价比价、采购订单、到货计划、价格和供应商绩效"),
            ("委外", "委外订单/发料/进度/对账", "委外工序、供应商产能、发料、回厂、质量、损耗和结算", "委外进度、收发料、回厂检验、损耗和委外对账"),
            ("行政", "人员/班次/考勤/劳务", "人员组织、考勤采集、异常、报表、劳务单位和结算规则", "考勤日报/月报、异常清单、工时汇总和劳务结算单"),
        ],
        [1200, 2100, 3600, 2460],
        font_size=7.9,
    )

    add_heading(doc, "财务部：全盘账、往来账、成本账与分析", 2)
    add_bullets(doc, [
        "负责总账、应收、应付、出纳、银行、资金计划、费用、固定资产、税务、预算、存货核算、制造成本和电子会计档案。",
        "往来账按客户、供应商、订单、发票、收付款、预收预付、账龄、到期日、争议和核销状态查询，自动生成客户/供应商对账单。",
        "全盘账覆盖凭证、明细账、总账、辅助账、科目余额、试算平衡、资产负债表、利润表、现金流量表和期末结转。",
        "成本账按产品、订单、工单、工序、批次、工作中心、设备和部门归集材料、人工、机时、委外、后工序、组装、返工、报废和制造费用。",
        "采购、委外、库存、生产、质量、行政劳务结算和销售业务按规则生成凭证并完成业务账、存货账、往来账、成本账和总账勾稽。",
        "财务工作台显示月结任务、对账差异、逾期应收、近期应付、资金缺口、成本异常、毛利变化、发票差异和审批超期。",
        "财务AI可分析往来、资金、成本、毛利、预算、月结、材料行情和异常原因；答案必须引用报表、凭证或源单据，不能直接记账、付款或关账。",
        "财务只能通过源单冲销、调整单或授权凭证处理差异，不能直接修改生产数量、检验结果、库存履历或行政考勤原始记录。",
    ])
    add_department_control_detail(doc, "财务部")

    add_heading(doc, "总经办：全厂经营分析、决策与重大预警", 2)
    add_bullets(doc, [
        "建立全厂驾驶舱，汇总订单、报价、销售、回款、毛利、现金、采购、委外、库存、在制、交付、质量、设备、人员和考勤。",
        "支持日、周、月、季度和年度视角，以及目标、预算、上期、同期和滚动预测对比；所有指标可下钻责任部门和源单。",
        "重点查看订单交付风险、关键缺料、瓶颈设备、供应商延期、跨序返工、报废、客户投诉、呆滞库存、逾期应收和资金缺口。",
        "重大报价、合同、订单变更、采购定标、委外、让步放行、报废、付款、预算和投资按金额或风险进入总经办审批。",
        "总经办AI支持“发生了什么、为什么、影响什么、建议关注什么”的连续追问和情景模拟，并明确事实、推断和预测。",
        "管理指令必须形成责任部门、责任人、截止时间、措施、验证和关闭，不能只通过口头或聊天工具下达。",
        "总经办拥有分析和审批范围不代表可以越权修改凭证、工艺、检验、库存和考勤；所有敏感钻取继续受字段权限约束。",
    ])
    add_department_control_detail(doc, "总经办")

    add_heading(doc, "生产部：CNC与车床制造执行", 2)
    add_bullets(doc, [
        "CNC和车床分别建立工作中心、设备组、班组、产能日历、标准工时、机时费率、技能要求和可加工产品/工序能力矩阵。",
        "PMC/APS下达计划后，生产主管按机台、人员、夹具、刀具、材料、程序和交期派工；现场人员扫码登录设备、工单和工序。",
        "开工前校验材料批次、图纸、工艺、CNC/车床程序版本、刀具、夹具、检具、人员资质、设备状态和上道工序。",
        "首次开机、换型、换刀、换程序、材料变更或返工后触发首件检验；首件未放行不得批量生产。",
        "报工记录投入、一次合格、不良、返工、报废、剩余、开始/结束、暂停、异常、人员、设备、程序和实际工时。",
        "CNC程序和车床程序由工程部发布受控版本，设备端只能调用当前有效程序；程序下载、修改、下发和使用均留日志。",
        "机台工作台显示今日任务、优先级、缺料、质量状态、计划/实际进度、下一工序和超期；班组看板显示产量、工时、良率、达成率和在制。",
        "设备记录运行、待机、停机、报警、换型、节拍、能耗、维修和OEE；停机原因必须从标准原因选择并支持补充说明。",
        "加工完成后按批次或单件码转序到品质、后工序、组装、委外或仓库，禁止无单、无检验和无扫码的线下流转。",
    ])
    add_department_control_detail(doc, "生产部（CNC与车床）")

    add_heading(doc, "品质部：首检、点检、末件检、FAI与质量档案", 2)
    add_bullets(doc, [
        "检验计划由产品、图纸版本、工艺、工序和客户要求自动带出，包含检验项目、气泡号、标准值、上下限、方法、检具、频率、抽样和关键等级。",
        "首检用于新批开工、换型、换程序、换刀、材料/设备变化和返工后验证；系统保存首件数量、全部测量值、照片、检具、人员、设备、程序和放行结论。",
        "点检/巡检按时间、数量、班次、风险或工序自动生成任务；到期未检、漏检、超差和重复异常进入预警，不能事后集中补录伪造过程记录。",
        "末件检在批次结束、换班、换型或工单关闭前对最后加工件执行，帮助发现加工漂移；若末件不合格，系统反查自最近一次合格检验后的影响数量。",
        "FAI用于新产品、打样、重大工程变更、客户要求或停产恢复；按气泡图逐项记录尺寸、材料、性能、外观和文件，系统生成受控FAI报告。",
        "来料、过程、完工、出货前全检和客户退货检验均存入QMS，记录批次/序列号、测量值、图片、缺陷、数量、结论、电子签名和时间。",
        "不合格自动冻结受影响数量，启动NCR/MRB、返工/返修、让步、报废、CAPA或8D；品质放行后才能转序、入库或出货。",
        "检具与量具记录编号、精度、量程、位置、校准证书和有效期；过期、停用或不适用检具不能用于提交检验结果。",
        "所有检验记录和FAI报告按客户、订单、产品、图纸版本、工单、工序、设备、人员、材料批次和检具长期检索并保留修改审计。",
    ])
    add_department_control_detail(doc, "品质部")

    add_heading(doc, "后工序部：清洗、表面标识、处理、全检包装与独立组装", 2)
    add_callout(doc, "组装必须独立", "组装具有独立BOM、齐套、领料、装配顺序、扭矩/参数、测试、序列号和成本，不能作为“包装备注”或普通后处理状态。后工序之间也必须按工艺路线逐站扫码，不能只在最后一次性报完。", fill="FFF4E5", color=GOLD)
    add_heading(doc, "清洗工作中心", 3)
    add_bullets(doc, [
        "工单维护清洗方式、槽液/清洗剂批次、浓度、温度、时间、漂洗、烘干、防锈、设备和装载量；扫码记录实际参数。",
        "清洗后检查油污、残屑、水渍、腐蚀和干燥状态，不合格返回清洗或进入质量异常。",
    ])
    add_heading(doc, "丝印工作中心", 3)
    add_bullets(doc, [
        "关联图稿/字体/位置/颜色版本、网版、油墨批次、调配比例、固化条件、样件确认和有效期，防止使用错版或过期油墨。",
        "检验文字、方向、颜色、附着力、清晰度、位置和外观；首件确认后才能批量丝印。",
    ])
    add_heading(doc, "镭雕工作中心", 3)
    add_bullets(doc, [
        "关联镭雕程序、内容模板、序列号规则、位置、字体、功率、速度和治具；程序版本由工程发布并校验产品匹配。",
        "支持单件码与镭雕内容绑定，检查漏雕、错码、重复码、位置和可读性，必要时使用视觉或扫码验证。",
    ])
    add_heading(doc, "打磨与振磨工作中心", 3)
    add_bullets(doc, [
        "打磨记录砂带/砂纸/磨料批次、粒度、工具、时间、人员和要求；检验毛刺、圆角、划伤、尺寸和表面一致性。",
        "振磨记录设备、介质、研磨液、配比、转速、时间、装载量和烘干；防止混料，并对外观、尺寸和残留介质检查。",
    ])
    add_heading(doc, "出货前全检与包装工作中心", 3)
    add_bullets(doc, [
        "根据客户、产品和订单生成100%全检任务，逐件/逐批扫码确认外观、关键尺寸、标识、数量、清洁、文件和质量状态。",
        "包装任务带出包装规范、数量/箱、内外标签、箱号、序列号、重量、保护材料、客户标签和拍照要求，生成装箱单和包装履历。",
        "全检不合格品隔离并关联返工/补货；未完成全检、标签校验、OQC和文件齐套时系统禁止发货。",
    ])
    add_heading(doc, "组装工作中心", 3)
    add_bullets(doc, [
        "按独立组装BOM和组装工单领料与齐套，记录组件批次/序列号、装配人员、工位、时间和一对一/一对多追溯关系。",
        "维护装配顺序、作业指导书、扭矩、压力、间隙、胶水/辅料批次、固化时间、专用工装和防错校验。",
        "装配过程可设置首件、过程检查、功能测试和最终测试；测试不合格生成拆解、返工、换件或报废流程。",
        "组装完工生成成品序列号、组件谱系、实际材料/人工/工时和组装成本，进入全检包装或成品入库。",
    ])
    add_department_control_detail(doc, "后工序部（含独立组装）")

    add_heading(doc, "工程部：BOM、程序、图纸、工艺与打样", 2)
    add_bullets(doc, [
        "建立物料工程属性、客户料号、图号、版本、原材料/成品尺寸、重量、材料、表面处理和单位换算。",
        "建立设计BOM、制造BOM和组装BOM，维护用量、损耗、替代料、投料工序、客户定制、生效范围和版本。",
        "集中存放公司图纸、客户图纸、3D模型、作业指导书、检验规范和技术协议，支持预览、水印、下载日志、生效和失效。",
        "CNC/车床编程程序按产品、工序、机型、刀具、夹具、版本和校验值管理；发布前完成编程、仿真、试切、首件和审批。",
        "程序由系统向授权设备或终端发布，现场修改必须回传并形成新版本，禁止长期使用个人电脑、U盘或聊天文件作为正式程序源。",
        "建立完整工艺路线，明确CNC、车床、清洗、丝印、镭雕、打磨、振磨、组装、全检包装和委外工序及检验点。",
        "打样使用样品订单/试制工单，记录材料、程序、工艺参数、问题、改进、实际工时、样品成本、首件和FAI结果。",
        "工程变更评估并联动BOM、程序、图纸、工艺、检验、库存、采购、在制、委外、报价和客户确认，不得只替换附件。",
    ])
    add_department_control_detail(doc, "工程部")

    add_heading(doc, "PMC：MRP、齐套、排程与生产进度", 2)
    add_bullets(doc, [
        "按销售订单、预测、样品、备料、安全库存和返工补料运行MRP，结合合格库存、已分配、冻结、不良、在制、采购和委外在途计算净需求。",
        "输出采购建议、生产建议、委外建议、调拨建议、缺料清单和动态齐套表；建议经责任人确认后转正式单据。",
        "建立周交付、周生产、日机台和后工序/组装计划，考虑CNC、车床、后工序、组装、品质、人员、工装和委外产能。",
        "跟踪订单、工单、各工序、采购、委外、品质和包装发货进度，显示计划/实际、当前工序、等待原因、责任人和预计完成。",
        "订单、库存、报工、报废、返工、质量冻结、采购到货、委外回厂和优先级变化后自动重算齐套、排程和交付风险。",
        "对缺料、资源冲突、插单、跨序返工、设备停机和供应商延期进行模拟，提交加班、替代设备、调整优先级或外发建议。",
        "PMC可以调整计划和优先级，但不能直接修改库存数量、检验结论、采购成交价和财务成本。",
    ])
    add_department_control_detail(doc, "PMC")

    add_heading(doc, "采购：询报价、金属价格、订单与到货", 2)
    add_bullets(doc, [
        "负责供应商准入与分类、采购询价、供应商网页报价、多轮议价、比价、定标、采购订单、交期确认、ASN、到货和退货协同。",
        "采购工作台显示MRP采购建议、紧急缺料、待询价、待定标、订单未确认、临期/超期到货、价格变动、来料不良和对账差异。",
        "金属材料采购引用交易所/持牌行情、供应商报价、历史成交、汇率、升贴水、加工和物流，保存下单价格快照并分析市场偏离。",
        "采购订单必须关联需求来源、物料/版本、数量、税率、币种、交期、质量、包装、付款和价格来源；变价、超量、提前或延期按规则审批。",
        "供应商ASN和送货单进入仓库与IQC，采购跟踪合格入库而不仅是物理到货；不合格、退货、补货和8D同步供应商绩效。",
        "采购对账基于采购订单、合格入库、退货、价格、税率和发票自动生成，采购确认业务差异，财务确认应付与付款。",
        "采购与委外岗位、订单类型、价格和对账分开；同一人员兼任时也必须通过角色和审批记录区分操作责任。",
    ])
    add_department_control_detail(doc, "采购")

    add_heading(doc, "委外：工序外发、供应商进度、回厂质量与结算", 2)
    add_bullets(doc, [
        "委外部门负责热处理、表面处理、特殊加工或产能外发等委外工序，维护委外供应商能力、工艺、价格、提前期、质量和产能。",
        "委外订单关联原生产工单、外发工序、图纸/工艺版本、外发数量、材料/半成品批次、质量要求、交期和后续返回节点。",
        "外发时扫描并记录发料/半成品、箱数、重量、批次、送货单和供应商签收；供应商在门户维护收料、开工、进度、待检、完成和发回。",
        "回厂扫描委外送货单，核对原发数量、合格、返工、报废、损耗和剩余；检验合格后才能回到指定内部工序。",
        "委外异常支持供应商返工、重新外发、退回内部、补料、索赔和8D，并同步订单交期、在制位置、质量成本和供应商绩效。",
        "委外对账按发料、供应商收料、合格回厂、损耗、返工、报废、扣款、加工费和运费生成，避免只按供应商发票付款。",
        "委外工作台和采购工作台可以共享供应商主数据，但委外进度、半成品在制、委外损耗与采购原料到货必须分别统计。",
    ])
    add_department_control_detail(doc, "委外")

    add_heading(doc, "行政：考勤采集、分析、报告与多劳务结算", 2)
    add_bullets(doc, [
        "维护员工、工号、部门、岗位、班组、用工类型、劳务单位、入离职、合同有效期、班次和考勤地点等基础数据。",
        "接入门禁/考勤机、移动打卡或经批准的文件导入，保存原始打卡时间、设备、地点、来源和导入批次；原始记录不可被覆盖。",
        "配置正常班、白夜班、跨日班、轮班、休息、法定节假日、调休、迟到早退、旷工、加班、请假、出差、培训和补卡规则。",
        "系统按日计算应出勤、实出勤、正常工时、加班、请假、迟到、早退、缺卡和异常；员工/主管在线确认或申诉并保留审批。",
        "生成个人、班组、部门、用工类型和劳务单位的考勤日报、周报、月报、异常清单、加班报告、请假报告和出勤趋势。",
        "维护多个劳务单位及不同结算规则，可按人数、出勤天数、正常/加班小时、班次、岗位、计件、管理费、补贴、扣款和税费计算。",
        "劳务结算以已审核考勤、MES有效报工/计件、合同费率和调整单为依据，生成劳务对账单、结算单和差异明细。",
        "劳务单位可在受控门户确认对账或由行政导入确认结果；行政审核人数与工时，部门确认用工，财务审核金额、发票、应付和付款。",
        "行政考勤与MES工时差异形成异常，例如已打卡无报工、已报工无打卡、跨部门借调和超长工时，供主管核实但不自动判定工资。",
        "考勤、身份证明、合同、工资和个人信息按严格机密管理；总经办看汇总，部门主管看本部门，非授权人员不能查看个人明细。",
    ])
    add_department_control_detail(doc, "行政")

    add_heading(doc, "跨部门数据交接与职责分离", 2)
    add_table(
        doc,
        ["上游", "交接对象", "下游", "系统控制"],
        [
            ("工程部", "物料、BOM、图纸、程序、工艺、检验版本", "PMC/生产/品质/采购/委外", "审批发布后生效；变更自动评估影响"),
            ("PMC", "采购/生产/委外建议、计划与优先级", "采购/委外/生产/后工序", "建议转正式单需确认；保留来源与变更"),
            ("采购", "采购订单、价格快照、ASN和到货计划", "仓库/品质/PMC/财务", "合格入库后更新供给与应付依据"),
            ("委外", "外发批次、进度、回厂数量和损耗", "PMC/品质/生产/财务", "半成品位置和数量全程可追溯"),
            ("生产/后工序", "报工、在制、异常、全检、包装与组装履历", "品质/PMC/仓库/财务", "扫码产生；禁止跨序、补录和无单移动"),
            ("品质", "检验、冻结、放行、返工和FAI", "全部业务部门", "品质结论控制转序、入库和发货"),
            ("行政", "审核考勤、工时和劳务结算", "部门主管/财务/总经办", "原始打卡、调整、审批和结算分层"),
            ("财务", "往来、成本、现金、预算和经营结果", "总经办/责任部门", "汇总可钻取源单；财务调整不反改业务履历"),
        ],
        [1300, 3100, 1700, 3260],
        font_size=8.2,
    )


def section_roadmap(doc):
    add_heading(doc, "分阶段实施路线", 1)
    add_para(doc, "系统建设采用“先规则与数据、再业务闭环、再精细计划与现场、最后外部协同和智能分析”的路线。PLM若是当前最大瓶颈，可与核心ERP并行试点；APS和设备联网必须建立在BOM、库存、工时和报工可信的基础上。")
    phases = [
        ("阶段0：准备与蓝图", "2-4周", "流程调研、范围、责任、数据和验收标准", "蓝图、RACI、实施计划、风险清单、验收脚本"),
        ("阶段1：基础数据与工程", "4-8周", "组织部门岗位、人员班次、物料、图纸、BOM、工艺、设备、工装、检验、库位和财务主数据", "编码规范、数据模板、版本规则、期初盘点、数据质量报告"),
        ("阶段2：经营与业财闭环", "2-4个月", "报价、合同、销售、采购、金属价格、委外、库存、基础工单、行政考勤、外贸、应收应付和自动凭证", "订单到回款、采购到付款、考勤到劳务结算、材料价格到报价成本、生产到成本闭环"),
        ("阶段3：计划、质量、成本与追溯", "2-5个月", "MRP、动态齐套、条码、批次、首检/巡检/末件/FAI、返工、实际成本、供应商进度和系统报表", "可解释MRP、正反追溯、工单成本、质量闭环"),
        ("阶段4：APS/MES/WMS与设备", "3-12个月", "有限排程、扫码执行、不跳工序、跨序返工、仓储作业、设备联网、OEE和预警中心", "现场实时、计划执行一致、异常闭环、仓库防错"),
        ("阶段5：门户、集团与经营分析", "持续建设", "供应商/客户门户、PLM深化、SRM、BI、集团财务、多工厂、套料/合批和AI辅助", "内外协同、经营驾驶舱、集团标准和持续优化"),
    ]
    add_table(doc, ["阶段", "参考周期", "建设重点", "主要交付"], phases, [1700, 1100, 3400, 3160], font_size=8.7)

    for title, objective, work, functions, acceptance in [
        ("阶段0：项目准备与业务蓝图", "明确建设边界、优先级、项目治理和可量化目标。",
         ["按财务、总经办、生产、品质、后工序、工程、PMC、采购、委外、行政以及业务/仓库职能逐部门访谈并现场走查。", "梳理纸质、Excel、微信群和重复录入点，确定未来流程与岗位责任。", "建立关键指标基线，包括订单准交率、库存准确率、一次合格率、返工率、报废率、月结周期、考勤准确率和报价周期。", "形成标准功能、增购模块、接口和定制清单，冻结首期范围。"],
         ["组织、角色、审批、编号、打印、移动端和接口总体设计。", "项目问题、变更、风险和决策台账。"],
         ["管理层批准范围与预算。", "各部门确认未来流程和数据责任。", "真实业务验收案例准备完成。"]),
        ("阶段1：基础数据与工程基线", "建立可支持MRP、成本、排程和追溯的唯一数据源。",
         ["清理公司、部门、岗位、班组、工作中心、人员、班次、劳务单位、物料、客户、供应商、仓库、科目和组织数据。", "完成关键产品BOM、组装BOM、CNC/车床程序、全工序路线、图纸、设备能力和检验标准。", "建立长宽高、直径、毛坯重量、成品重量、材料、损耗和余料属性。", "进行上线前实物盘点和期初余额核对。"],
         ["主数据审批、重复校验、生效失效、版本、批量导入和修改日志。", "PLM/ERP之间的物料、BOM和工艺发布规则。"],
         ["关键BOM完整率达到项目约定门槛。", "库存账实准确并能定位到库位/批次。", "图纸、工艺和检验标准版本一致。"]),
        ("阶段2：核心业务与财务闭环", "让客户订单、采购、生产、交付和财务凭证在系统中完整流转。",
         ["上线报价、成本分析、合同评审、订单、采购、金属材料价格中心、委外、库存、基础工单、销售出库、发票和收付款。", "配置样品、批量、备料、阶梯价格、客户信用、供应商价格、行情/牌号映射、落地价和委外结算。", "建立外贸商业发票、装箱单和报关数据包。"],
         ["订单到回款、采购到付款、工单领料到完工入库。", "业务单据按规则自动生成财务凭证。"],
         ["一张真实订单完成全链路演示。", "库存、业务、往来和总账可对账。", "客户正式报价PDF由系统生成。", "采购、财务和业务按权限看到同一材料价格快照及不同明细。"]),
        ("阶段3：精细计划、质量、成本和追溯", "提高齐套、计划、质量和成本数据的准确性。",
         ["校准提前期、安全库存、批量、损耗、在途和在制，运行MRP。", "上线条码、批次、首检、点检/巡检、末件检、FAI、IQC/FQC/OQC、内部返工、委外质量和客户退货。", "建立标准成本、实际成本、后工序、组装、返工和报废成本。", "上线系统生成的订单、工序、采购、供应商、齐套、考勤、劳务结算、检验、对账和分析报表。"],
         ["动态净需求、齐套率、缺料清单和建议采购/生产量。", "正向与反向追溯、NCR/MRB/CAPA。"],
         ["MRP每条建议可追溯需求来源。", "第五工序退回第二工序后直接返回第五工序，以及从第三、第四工序顺序重流两种跨序返工场景均跑通。", "报价、标准和实际成本可比较。"]),
        ("阶段4：智能工厂执行", "实现计划下达到现场执行、仓储和设备的实时闭环。",
         ["配置APS约束、日/班次计划、插单模拟和产能负荷。", "上线CNC、车床、清洗、丝印、镭雕、打磨、振磨、全检包装和独立组装的MES扫码执行。", "上线不跳序、返工、异常、人员设备绩效，以及WMS ASN收货、上架、拣配、余料、盘点和仓位图。", "建设统一预警和设备维修/OEE。"],
         ["计划、执行、库存和成本之间接口。", "现场看板、Andon、预警升级和异常知识库。"],
         ["工序进度实时且来自扫码/采集。", "超期工序和审批自动预警。", "仓库关键作业有条码防错。"]),
        ("阶段5：外部协同与经营优化", "让客户、供应商、集团和管理层在同一数据链上协同。",
         ["供应商在线报价、议价、订单确认、ASN、送货单、质量、对账、发票和付款查询。", "客户查看受控项目里程碑、交货计划、发货、质量文件和返工进度。", "深化PLM、SRM、多工厂、集团财务、BI和绩效。", "实施余料匹配、套料、订单合批，以及面向财务和老板的受控AI分析问答。"],
         ["门户数据隔离、外部账号、文件水印和下载日志。", "经营驾驶舱、指标语义层、AI只读问答、引用和审计。"],
         ["供应商报价到中标订单、送货和对账全程在线。", "客户只能看到授权项目和文件。", "BI指标与业务、财务明细可钻取对账。", "AI固定问题集与人工报表数字一致且不能越权。"]),
    ]:
        add_heading(doc, title, 2)
        add_para(doc, f"阶段目标：{objective}", bold_prefix="阶段目标：")
        add_heading(doc, "主要工作", 3)
        add_bullets(doc, work)
        add_heading(doc, "需要上线的能力", 3)
        add_bullets(doc, functions)
        add_heading(doc, "验收门槛", 3)
        add_bullets(doc, acceptance)


def section_master_engineering(doc):
    add_heading(doc, "主数据、产品工程与PLM", 1)
    add_heading(doc, "物料与产品主数据", 2)
    add_para(doc, "物料主数据是MRP、库存、报价、质量和成本的共同基础。尺寸不得全部写入备注，应建立可计算、可检索的结构化字段。")
    add_bullets(doc, [
        "编码、名称、客户料号、图号、版本、物料分类、采购件/自制件/委外件/客供件属性。",
        "材料牌号、形态（板、棒、管、锻件、铸件）、原材料长宽高或直径、下料尺寸、成品长宽高或直径。",
        "毛坯重量、成品重量、密度、基本单位、采购单位、生产单位、库存单位和换算关系。",
        "采购提前期、生产提前期、委外提前期、安全库存、最小采购量、采购倍数、标准损耗和保质期。",
        "批次、炉号、序列号、质保书、检验方式、存储条件、默认库位、默认供应商和成本方法。",
    ])
    add_heading(doc, "BOM与组件关系", 2)
    add_bullets(doc, [
        "支持多层BOM、设计BOM、制造BOM、虚拟件、替代料、可选件、客户定制BOM和工程试制BOM。",
        "每个组件维护单位用量、损耗、投料工序、是否倒冲、是否客供、是否批次追溯和余料复用规则。",
        "支持反查某物料用于哪些产品，并在工程变更时分析库存、采购、委外和在制影响。",
        "BOM必须有版本、生效日期、失效日期、审批人和适用订单范围，禁止直接覆盖历史版本。",
    ])
    add_heading(doc, "图档与工程版本", 2)
    add_bullets(doc, [
        "上传PDF、DWG、DXF、STEP、IGES、Excel技术要求、图片、样品照片、检验标准和客户邮件附件。",
        "支持在线预览、格式转换、版本对比、水印、下载记录、权限、生效/失效和与报价、订单、产品、工单关联。",
        "图纸变更触发工程变更单，自动评估已采购、已领料、在制、委外、库存和交付影响。",
        "图纸自动识别尺寸、材料、公差和表面处理属于高级能力，应作为PLM/CAD接口或专项定制单独验收。",
    ])
    add_heading(doc, "工艺路线与制造资源", 2)
    add_bullets(doc, [
        "维护工序顺序、工作中心、可用机台、人员技能、装夹次数、夹位、准备时间、加工时间、转移批量和等待时间。",
        "维护工序材料消耗、标准良率、标准损耗、是否委外、是否关键工序、是否需要首检/巡检/全检。",
        "关联程序版本、作业指导书、图纸、刀具、夹治具、模具、检具、包装要求和安全要求。",
        "图纸修改后必须评估并更新工艺路线；正在执行的工单通过临时工艺变更形成新版本，不得直接覆盖。",
    ])
    add_heading(doc, "工装、夹具、模具、刀具与检具", 2)
    add_bullets(doc, [
        "建立编号、名称、图片、供应商、适用产品/工序/设备、存放位置、数量、状态和负责人。",
        "记录使用次数、寿命、保养、维修、校验、停用、替代和摊销成本。",
        "排程时将关键夹具、模具和检具作为有限资源，避免多工单同时占用。",
        "检具过期、停用或精度不满足时不得被检验任务选用。",
    ])
    add_heading(doc, "设备、人员与能力矩阵", 2)
    add_bullets(doc, [
        "每台设备关联可生产产品、可执行工序、加工范围、标准节拍、每日产能、班次、维护日历和替代设备。",
        "每个操作人员和技术员维护可操作设备、技能等级、资格证书、有效期和允许工序。",
        "派工前校验人员、设备、图纸、程序、工装、材料和检验准备状态。",
    ])
    add_heading(doc, "主数据治理", 2)
    add_bullets(doc, [
        "明确物料、BOM、工艺、客户、供应商、设备、工装、检验、仓位和财务主数据责任人。",
        "新增、修改、停用、合并和批量导入均需审批与日志；重复编码和关键字段缺失应阻止提交。",
        "定期发布数据质量报表：重复率、完整率、失效数据、孤立BOM、无工艺产品、负库存和单位换算异常。",
    ])


def section_sales_customer(doc):
    add_heading(doc, "业务报价、合同、订单与客户协同", 1)
    add_heading(doc, "客户询价与报价分类", 2)
    add_bullets(doc, [
        "询价记录客户、图纸、技术要求、数量、样品/批量/备料分类、期望交期、包装、运输、税率和贸易条件。",
        "每次报价保留编号、版本、状态、有效期、未成交原因和转订单关系。",
        "支持样品价、小批量价、批量价、年度框架价和阶梯数量价格。",
    ])
    add_heading(doc, "业务订单类型与收费方式", 2)
    add_bullets(doc, [
        "模具订单：客户可以付费，也可以免费或部分收费；订单必须明确模具所有权、费用承担、预计总成本、押金/返还、正式订单分摊和未回收成本。",
        "样品订单：客户可以付费，也可以免费或部分收费；必须关联报价/商机、样品目的、客户确认、预计转正式订单日期和费用承担部门。",
        "正式业务订单：客户必须付费；允许预收、账期、分期、尾款和质保金，但单价或总额为零、未设置付款条件、缺核价或缺原始客户订单时禁止提交。",
        "订单类型和收费方式分别维护；免费不免成本核算，三类订单均需归集材料、人工、机时、委外、检验、返工、报废、包装和物流实际成本。",
        "模具/样品转正式订单通过来源关系关联，不覆盖原单；抵扣、返还或成本分摊保留逐笔履历。",
    ])
    add_heading(doc, "香港代生产客户正式订单价格", 2)
    add_bullets(doc, [
        "客户主数据增加“香港代生产价格客户”勾选项；香港代生产不是新的订单类型，也不要求业务在订单上另选特殊模式。",
        "业务录入客户订单上的原始价格；只有客户已勾选且订单类型为正式业务订单时，系统自动生成原始价格×70%的正式订单价格。",
        "样品订单和模具订单即使选择该客户也保持原输入价格；未勾选客户的正式、样品和模具订单均不执行70%转换。",
        "系统保存客户勾选快照、订单类型、原始价格、70%系数、计算后价格和版本；除价格换算外，订单审核和全部下游流程与普通订单完全相同。",
    ])
    add_heading(doc, "历史与相似产品报价", 2)
    add_bullets(doc, [
        "按客户、产品、图号、材料、原材料尺寸、成品尺寸、重量、工艺、机台、表面处理、数量和日期查询。",
        "展示历史报价、成交价、最后采购价、委外价、报价成本、实际成本和毛利偏差。",
        "建立材料＋工艺＋尺寸范围＋精度等级等产品特征标签，支持相似产品推荐；图形相似度属于专项能力。",
    ])
    add_heading(doc, "机加工成本分析与报价模型", 2)
    add_bullets(doc, [
        "材料：牌号、毛坯尺寸/重量、成品重量、实时/日结行情、供应商/历史采购参考、价格快照、汇率、单价、利用率、标准损耗、余料和废料回收价值。",
        "加工：工序、设备、装夹次数、准备时间、单件工时、机时费率、人工费率、刀具消耗和工装摊销。",
        "委外：工序、供应商历史价、最低批量费、运输、损耗、返工责任和交期。",
        "其他：检验、包装、物流、管理费用、财务费用、风险系数和目标毛利。",
        "数量阶梯应自动分摊编程、调机、模具、最低采购量和委外最低批量费用。",
    ])
    add_heading(doc, "报价审批与正式文件", 2)
    add_bullets(doc, [
        "根据毛利、折扣、总额、新材料、新工艺、紧急交期、客户信用和工装投资触发分级审批。",
        "工程核价、采购外发价、财务毛利和管理层授权未完成时不得生成正式客户版。",
        "审批后锁定版本，系统生成中文/英文PDF；客户版隐藏内部成本、供应商价格和毛利。",
        "材料行情变化超过企业阈值、价格快照过期或汇率明显变化时，报价进入重新核价待办；未重新审批不得覆盖原报价版本。",
        "报价转订单不重复录入；图纸、数量或交期重大变化时提示重新核价。",
    ])
    add_heading(doc, "合同评审和订单控制", 2)
    add_bullets(doc, [
        "合同扫描件、技术协议、质量协议、付款、交期、违约、纸质档案位置和财务保管责任人统一归档。",
        "批量订单下达前由业务、工程、PMC、采购、品质、财务和管理层完成合同评审。",
        "未通过评审时禁止下批量工单、采购订单、委外订单和新工装。",
        "订单变更自动分析材料、在制、委外、库存、呆滞、成本和其他客户交期影响。",
    ])
    add_heading(doc, "客户退货、补货和售后", 2)
    add_bullets(doc, [
        "客户质量问题关联原销售订单、出货批次/序列号、图纸版本、检验记录和材料批次。",
        "支持退货、现场筛选、返修、补货、让步使用和客户8D，形成RMA及质量异常。",
        "客户退货返工生成独立返工工单和路线，记录材料、人工、机时、委外、复检、加急运输和毛利损失。",
    ])
    add_heading(doc, "客户协同门户", 2)
    add_bullets(doc, [
        "客户使用网页查看自己的报价、合同、订单、项目里程碑、交货批次、预计交期、发货和质量文件。",
        "客户可确认报价、图纸版本、样品、交货计划和收货，可提交客诉、退货和补货需求。",
        "客户可下载报价单、订单确认书、交货计划、发货通知、送货单、装箱单、商业发票、检验报告、材质证明和8D。",
        "客户可见进度采用受控里程碑：技术评审、材料准备、生产、委外、检验、包装、待发货、已发货、已签收。",
        "禁止默认展示内部成本、供应商名称与价格、操作人员、内部责任认定和其他客户信息。",
    ])
    add_heading(doc, "外贸业务与报关资料", 2)
    add_bullets(doc, [
        "维护客户英文名称、地址、收货人、通知人、HS编码、英文品名、材质、用途、品牌、原产国、贸易条款和运输信息。",
        "发货审核后按实际数量、箱数、毛重、净重、币种和价格生成形式发票、商业发票、装箱单、销售合同和报关数据包。",
        "正式报关申报可通过接口连接单一窗口、报关行或货代系统；系统保留申报版本和回执。",
        "香港代生产价格客户的正式订单仍沿用相同的外贸与报关流程；客户勾选只触发订单价格70%换算，不增加或减少报关步骤。",
        "包装实绩或订单价格变化后，系统比较订单、合同、商业发票、装箱单、运输和申报数据；已提交版本只能更正或生成新版本，不能直接覆盖。",
    ])


def section_procurement_supplier(doc):
    add_heading(doc, "采购、委外与供应商协同", 1)
    add_heading(doc, "采购与委外业务闭环", 2)
    add_bullets(doc, [
        "采购申请可来自MRP、库存补充、工程试制、设备维修或人工紧急需求。",
        "询价、比价、定标、采购订单、供应商确认、ASN、到货、IQC、入库、退货、发票、应付和付款形成闭环。",
        "委外支持发料、供应商收料、过程进度、回厂、检验、返工、损耗、加工费和对账。",
    ])
    add_heading(doc, "供应商在线询价与报价", 2)
    add_para(doc, "供应商报价必须在网页中结构化填写，不能以下载Excel、邮件回传、采购代录作为标准流程。")
    add_bullets(doc, [
        "询价类型：原材料、标准件、定制零件、委外工序、模具工装、设备备件、紧急询价和年度框架询价。",
        "询价内容：物料/图号/版本、数量阶梯、技术和质量要求、包装、交期、交货地、币种、税率、贸易条件、付款和截止时间。",
        "供应商填写含/未税单价、运费、包装费、模具费、打样费、最低订单、阶梯价、产能、交期、报价有效期和技术偏离。",
        "支持元/件、公斤、米、平方米、套、批、小时、炉、模和工序等报价单位及换算。",
        "支持多轮议价、V1/V2/V3版本、密封报价、截止控制和报价解密审计。",
        "图纸带供应商水印，记录预览、下载和版本；供应商不能看到其他供应商、目标价、客户价或内部成本。",
    ])
    add_heading(doc, "实时金属材料价格中心：行情与企业价格体系", 2)
    add_bullets(doc, [
        "统一管理铜、铝、钢材、不锈钢、镍、锌、锡、铅、钛及企业自定义合金等材料的市场基准价、供应商报价、历史采购价和企业结算价。",
        "外部基准行情可接入上海期货交易所、伦敦金属交易所或企业购买的持牌数据商；具体品种、实时性、使用方式和展示范围以数据许可为准。",
        "行情记录至少包含数据源、交易所、品种/合约、牌号映射、现货/期货、交割月、买卖/结算价、币种、单位、时区、行情时间、接收时间和质量状态。",
        "供应商价格包含供应商、材料牌号、形态、规格区间、产地/品牌、税率、币种、数量阶梯、加工费、升贴水、运费、付款、交货地、有效期和最低订单量。",
        "内部采购历史价关联采购订单、供应商、到货批次、实际重量、含税/未税、运费、质量扣款和发票结算，形成企业真实成交价格曲线。",
        "建立材料牌号与交易所品种的受控映射；无法直接对应的合金、板棒管材和特殊规格，通过基准金属含量、升贴水、加工费或供应商指数形成企业参考价。",
        "企业采购落地参考价 = 市场基准价×汇率×单位换算 + 升贴水 + 加工/切割费 + 包装运输保险 + 关税及不可抵扣税费；每一项参数均可追溯。",
        "价格按实时、延时、日结或人工审批更新，页面必须显示“实时/延时分钟数/上一交易日/人工价”和最后更新时间，禁止把过期行情显示为实时。",
    ])
    add_heading(doc, "实时金属材料价格中心：采购、财务与业务视图", 2)
    add_bullets(doc, [
        "采购视图：查看市场基准、供应商最新报价、历史成交、价差百分比、数量阶梯、有效期和价格趋势，用于询价、比价、锁价和采购时点判断。",
        "采购下单时保存行情快照、汇率、升贴水和供应商报价版本；市场价格变化不能自动改写已审批采购订单，变价必须走订单变更审批。",
        "财务视图：查看价格变动对标准成本、暂估、在途采购、库存价值、在制成本、采购承诺、现金需求和订单毛利的模拟影响。",
        "财务成本重估仅生成模拟版本和差异报告，不直接修改法定存货账；正式标准成本调整、减值或会计处理须按会计政策审批。",
        "业务视图：在客户报价时查看经授权的材料参考价、趋势、成本影响、建议报价有效期和价格联动条款，不展示供应商身份、底价、采购议价记录或未授权财务信息。",
        "报价版本保存所用材料价格快照、行情时间、汇率、损耗率和有效期；行情超过阈值或报价即将到期时提醒业务重新核价。",
        "角色可见范围由材料品类、组织、字段和用途控制；采购可见供应商明细，财务可见成本与价值，业务默认只见经批准的参考价和报价影响。",
        "财务和老板AI可询问材料价格涨跌、供应商报价偏离、采购机会、库存与订单毛利影响，但答案必须注明行情源、时间、许可状态、单位和模拟假设。",
    ])
    add_heading(doc, "实时金属材料价格中心：预警、许可与数据质量", 2)
    add_bullets(doc, [
        "按材料、品种和角色配置日内/日间涨跌幅、与供应商报价偏差、报价成本失效、采购订单变价和数据中断预警。",
        "价格预警进入统一待办中心，可通知采购、财务、业务和管理层；责任人记录是否重询价、重报价、锁价、暂缓采购或接受风险。",
        "行情接口具备认证、限流、重试、幂等、断点续传、时区转换、交易日历、异常值校验、备用源和数据质量评分。",
        "源数据缺失、延迟超标、币种或单位异常、价格跳变时标记不可用于自动核价，并回退到最近有效快照或人工审批价。",
        "系统保存原始行情、标准化结果、换算公式、人工修正、审批、使用单据和失效状态，确保报价和成本可按历史时点重现。",
        "上线前由法务和采购确认交易所/数据商的展示、内部使用、估值、定价、派生数据、缓存、用户数和再分发许可；未经许可不得抓取网页或向客户/供应商传播行情。",
    ])
    add_heading(doc, "自动比价、评审和定标", 2)
    add_bullets(doc, [
        "系统比较含税/未税单价、运费、模具摊销、最低批量、样品/批量交期、月产能、付款、技术偏离、历史准交率和合格率。",
        "综合采购成本考虑单价、运输、包装、工装、检验、资金、预期质量损失和交付风险。",
        "采购、工程、品质、PMC、财务和管理层分别评审；支持单一中标、多源比例、主备供应商和部分数量中标。",
        "定标后自动转采购/委外订单；对中标价格、供应商或付款条件的再次修改必须重新审批。",
    ])
    add_heading(doc, "供应商订单、进度和交期协同", 2)
    add_bullets(doc, [
        "供应商网页登录查看自己的采购/委外订单、图纸版本、数量、已交、未交、交期、包装和检验要求。",
        "供应商可接受、拒绝、建议交期、拆分交货批次并维护已接单、已收料、加工中、待检、已完成、已发货等状态。",
        "供应商承诺交期和进度实时回写采购进度、委外进度、PMC齐套和客户交付风险。",
        "延期必须填写原因、完成数量、剩余数量、新预计交期和改善措施。",
    ])
    add_heading(doc, "ASN、送货单和标签", 2)
    add_numbered(doc, [
        "供应商从已确认订单选择本次送货明细，填写数量、批次、箱数、毛净重、炉号、质保书、车牌、司机、物流单号和预计到达时间。",
        "系统校验订单状态、剩余未交量、提前/超量送货规则、图纸版本、质量文件和到货预约。",
        "校验通过后生成ASN预到货编号、A4送货单PDF、箱标签、托盘标签和单件/批次二维码。",
        "仓库扫描送货单二维码调出订单并收货，随后完成IQC、入库或退货，采购和PMC进度自动更新。",
    ])
    add_heading(doc, "供应商质量协同", 2)
    add_bullets(doc, [
        "来料或委外异常向供应商发布不良批次、数量、缺陷、检验数据、图片、责任要求和回复期限。",
        "供应商可确认责任、提出异议、提交8D、补货、返工、退货和纠正预防措施。",
        "系统跟踪首次响应、原因分析、整改、补货、验证和关闭时长，并进入供应商质量绩效。",
    ])
    add_heading(doc, "采购与委外对账", 2)
    add_bullets(doc, [
        "采购对账按采购订单、合格入库、退货、价格、税率和供应商发票三方匹配生成。",
        "委外对账按委外订单、发料、供应商收料、合格回厂、返工、报废、损耗扣款、加工费和运输费生成。",
        "供应商门户在线确认、提出差异、上传发票并查看发票审核、计划付款和已付款状态。",
    ])


def section_planning(doc):
    add_heading(doc, "PMC、MRP与APS计划管理", 1)
    add_heading(doc, "需求与供给口径", 2)
    add_bullets(doc, [
        "需求包括销售订单、预测、样品、备料、安全库存、返工补料、工程试制和替换需求。",
        "供给包括合格库存、可用余料、在制预计完工、采购在途、委外在途、工厂间调拨和可替代料。",
        "待检、冻结、不良、已分配、客户专用和已预留库存必须从可用量中分离。",
    ])
    add_heading(doc, "MRP净需求计算", 2)
    add_callout(doc, "核心公式", "当前可用库存 = 合格库存 - 已分配量 - 冻结量；按期可供量 = 当前可用库存 + 交期前可完工在制 + 交期前可到货采购 + 交期前可回厂委外；净需求 = 毛需求 + 安全库存 + 预计损耗 - 按期可供量；建议采购/生产量再按最小批量和采购倍数取整。")
    add_bullets(doc, [
        "MRP展开销售订单和BOM，综合库存、在制、在途、损耗、提前期、最小批量、替代料和供应商日历。",
        "输出采购建议、生产建议、委外建议、调拨建议、缺料表、齐套表、预计到料和预计完工日期。",
        "系统自动生成的是采购申请或计划订单，未经审批不得直接成为对外采购订单。",
        "每一条建议必须能追溯来源订单、BOM层级、需求日期、现有供给和计算参数。",
    ])
    add_heading(doc, "动态成品备料与齐套", 2)
    add_bullets(doc, [
        "按销售订单和成品展开全部组件，显示单位用量、损耗、总需求、合格库存、已分配、待检/冻结、在制、采购/委外在途和缺口。",
        "订单、库存、报工、报废、质量冻结、采购到货、委外回厂、交期或优先级变化后实时重算。",
        "同时显示数量齐套率、物料种类齐套率、关键物料状态、最晚齐套日期和责任采购员/供应商。",
        "关键物料短缺时，即使普通物料齐全也不得显示完全齐套。",
    ])
    add_heading(doc, "周计划、日计划和交付优先级", 2)
    add_bullets(doc, [
        "计划层级包括月度主计划、周交付计划、周生产计划、日机台计划、班组计划和个人任务。",
        "日计划包含机台、人员、产品、工单、工序、数量、标准工时、计划开始/结束、材料、图纸、程序、工装和检验准备状态。",
        "优先级考虑客户等级、合同交期、延期风险、缺料、返工补货、后续装配、可合批性和管理层指令。",
        "人工调整优先级记录调整人、原因、时间和对其他订单的交期影响。",
    ])
    add_heading(doc, "APS资源约束与模拟", 2)
    add_bullets(doc, [
        "约束包括设备加工能力、人员技能、夹具/模具/检具、物料齐套、班次、维护、前后工序、委外周期、换型和最小批量。",
        "提供有限产能甘特图、设备负荷、订单交期模拟、插单模拟、多方案比较、缺料冲突和产能冲突。",
        "系统建议可替代设备、调整优先级、外发、加班或拆批方案，计划员确认后下发MES。",
    ])
    add_heading(doc, "订单合批、套料与余料优化", 2)
    add_bullets(doc, [
        "识别相同材料、规格、工艺、机台、刀具和表面处理的订单，提示合批生产。",
        "结合原材料长宽高、棒材长度、余料尺寸和订单交期进行套料或下料优化。",
        "比较合批前后材料利用率、换型/开机次数、制造周期、在制和交期风险。",
        "该能力通常需要APS、套料软件或定制算法，必须采用真实订单和余料数据专项验收。",
    ])
    add_heading(doc, "架机表与机台负荷管理", 2)
    add_para(doc, "架机表＝按机台×时间段的上机计划：哪台机、什么时间、装哪个工单/产品/工序、用什么夹具和程序、计划产出多少；是APS排程结果面向车间的落地视图，也是操作工傻瓜式任务清单的来源。")
    add_bullets(doc, [
        "MRP确认后系统自动生成初版架机表，输入包括工艺路线可用机台、夹位数、每夹位加工时间、准备/调机时间、每台机日产能、班次日历、设备维护计划、人员技能矩阵和在制占用。",
        "生成逻辑：优先满足交期，其次减少换型/开机次数（与订单合批联动），再平衡各机台负荷；初版为建议状态，未经PMC确认不下发MES。",
        "生成和调整时校验机台、夹具、模具、检具、程序和技术员共用冲突，冲突项标红并给出可替代方案。",
        "PMC可调整机台、顺序和时间；每次调整记录调整人、原因、时间和受影响订单交期模拟；插单、返工子订单和优先级变化触发重算建议。",
        "架机率＝机台已排产时间/可用时间，按机台、班组、日/周/月统计；架机计划与MES实际扫码开工对比，计划未上机、上机晚点、提前拆机进入偏差分析。",
        "操作工不能自行改变上机顺序；预警覆盖机台负荷超上限、关键订单未排入、架机冲突未处理、计划已到未上机。",
        "报表：初版/当前架机表（按机台和按订单两个视角）、架机率日报/周报、计划与实际上机偏差表、机台负荷平衡分析。",
    ])
    add_heading(doc, "平衡化生产", 2)
    add_bullets(doc, [
        "周/日计划层面显示各机台和工作中心负荷率热力图，识别瓶颈和闲置。",
        "系统建议拆批、转移可替代机台、调班或外发以拉平负荷，采纳建议同样调整留痕。",
        "目标：瓶颈工序满载、非瓶颈按节拍拉动、在制总量受控，并与长时间工单未完结监控联动。",
    ])
    add_heading(doc, "计划责任与进度看板", 2)
    add_bullets(doc, [
        "按订单显示图纸、工艺、材料、采购、委外、工装、刀具、程序、检验、生产、包装和发货状态。",
        "每个状态有责任人、计划时间、实际时间、异常原因、预计恢复和升级对象。",
        "提供局部/整体齐套、采购、委外、供应商、各工序、每张工单和客户交货进度。",
    ])


def section_mes(doc):
    add_heading(doc, "生产执行、设备与MES", 1)
    add_heading(doc, "工单与派工", 2)
    add_bullets(doc, [
        "生产订单可按交货批次、工艺路线、设备能力和转移批量拆分工单。",
        "派工前校验材料、图纸、程序、工装、刀具、检具、设备状态、人员资质和上道工序。",
        "任务显示计划数量、标准工时、优先级、交期、当前缺料、质量状态和电子作业指导书。",
    ])
    add_heading(doc, "一码到底与扫码执行", 2)
    add_bullets(doc, [
        "建立原材料批次码、工单码、流转批次码和单件序列码；精密件采用一物一码，普通批量件可一批一码。",
        "二维码关联客户订单、产品、图纸/工艺版本、材料批次、工序、设备、人员、时间、检验、返工、入库和发货。",
        "每道工序按人员码、设备码、工单/产品码顺序校验并扫码开工、暂停、异常、完工和转序。",
        "报工记录投入、一次合格、不良、返工、报废和剩余数量，不能只填一个完成百分比。",
    ])
    add_heading(doc, "工序成本事实采集", 2)
    add_bullets(doc, [
        "每道工序开工时生成唯一工序执行实例并关联成本卡；扫码、报工、设备采集、领退料、刀具工装、检验和数量交接共同形成成本事实。",
        "准备/调机、批量加工、人工、设备运行、停机、等待和检验时间分段记录，避免全部计入一个总工时。",
        "CNC、车床、切料、打磨、清洗、振磨、镭雕、丝印、组装、全检和包装分别报工；同名工序重复执行也分别编号和计费。",
        "现场只记录数量、时间、资源、材料和异常，不显示或允许修改敏感费率；财务按有效费率和政策计算成本。",
        "未报人员/设备/材料/数量/时长的工序不能完工或进入成本结算；离线补传必须保留实际发生时间和补录审批。",
    ])
    add_heading(doc, "工序防跳与傻瓜式操作", 2)
    add_bullets(doc, [
        "上道未完成、首件未合格、关键检验未完成、图纸版本不一致或质量冻结时，系统禁止下一工序。",
        "操作界面仅显示当前人员允许执行的任务、图纸、程序、数量和操作按钮，减少自由录入。",
        "有条件跳工序必须预先配置适用产品、条件、审批人、原因和后续补检要求，并保留日志。",
        "条码、设备、程序、材料和工装不匹配时系统即时提示并阻止开工。",
    ])
    add_heading(doc, "中途改工序控制", 2)
    add_bullets(doc, [
        "批次生产中途需要增加、删除、调整顺序或改换工序（中途改工序）时，必须通过工程变更生成新工艺路线版本，限定适用的订单、批次和数量，禁止现场口头改道。",
        "改图纸后必须同步修改工艺路线；新版本发布后系统自动重算受影响批次的交期、架机、检验计划和工序成本，未走版本变更的批次仍按原路线防跳校验。",
        "已完成工序的履历保持不变；变更点之后按新路线派工、报工和检验，新旧路线执行履历在同一批次谱系下可对照追溯。",
        "中途改工序需工程、品质、PMC会签，涉及客户要求或关键特性时增加业务确认；全部变更留审批和生效日志。",
    ])
    add_heading(doc, "内部质量返工与跨工序退回", 2)
    add_para(doc, "返工不仅包含客户退货，还包括来料、首件、自检、巡检、下工序、完工和出货检验发现的内部质量问题。")
    add_bullets(doc, [
        "同工序返工：质量异常判定后生成返工任务，独立记录返工工时、设备、刀具、材料和复检。",
        "跨工序退回：例如第五工序发现第二工序问题，系统冻结受影响批次和数量，由质量、工程、生产共同判定影响特性、需重做工序和后续路线。",
        "路线A—返工后返回发现工序：第五工序退回第二工序返工，第二工序复检合格后直接返回第五工序继续加工；第三、第四工序原合格结果继续有效，但必须保留返工评审依据和有效性确认记录。",
        "路线B—返工后顺序重新流转：第五工序退回第二工序返工后，因第二工序加工改变了第三、第四工序的基准、尺寸、表面或装配状态，不能直接返回第五工序；系统必须按第二→第三→第四→第五的顺序重新派工、报工和检验。",
        "路线判定字段至少包括返工起点、缺陷来源、受影响特性、原后续工序是否仍有效、必须重做工序、必须复检项目、回流策略、目标返回节点和批准人。",
        "同一批次可以按数量拆分路线：部分数量走路线A直接返回第五工序，另一部分走路线B从第三、第四工序重新流转；系统生成子批次、独立二维码和独立在制数量，禁止混批。",
        "路线B执行时，第三、第四工序原履历不得删除，应标记为被返工路线替代；重新执行产生新的工序履历、版本、人员、设备、工时、检验和成本。",
        "返工路线定义返工起点、重做工序、可保留工序、复检项目、顺序约束、返回节点和重新汇合条件；未完成规定复检和顺序工序不得回到第五工序。",
        "临时增加或修改工序通过工程变更生成新路线版本，限定适用批次和数量，并重算交期与成本。",
        "返工路线批准后同步重算各工序在制、机台负荷、人员任务、物料补领、预计完工、客户交期、返工成本、一次合格率和最终良率，并触发相关预警。",
        "返工完成后由质量复检和放行；操作人员只能按系统当前返工任务扫码执行，不能自行选择返回节点、跳过第三/第四工序或解除质量冻结。",
    ])
    add_heading(doc, "设备、人员与绩效", 2)
    add_bullets(doc, [
        "按机台和人员统计每日任务、投入、良品、不良、返工、报废、标准工时、实际工时和达成率。",
        "设备统计运行、停机、待机、报警、节拍、产量、良率、换型和OEE；人员统计一次合格率和效率。",
        "计时/计件工资从已审核报工和工价规则生成，返工、报废和异常责任处理按制度配置。",
    ])
    add_heading(doc, "设备台账、点检、保养与维修", 2)
    add_bullets(doc, [
        "设备台账包含部门、规格、能力、投产日期、关键备件、维护日历和责任人。",
        "支持点检、保养、故障报修、维修申请、备件领用、验收、停机原因和维修成本。",
        "维修、保养和故障状态同步APS，避免向不可用设备排程。",
    ])
    add_heading(doc, "现场异常闭环", 2)
    add_bullets(doc, [
        "异常类型包括缺料、设备、工艺、图纸、程序、质量、人员、模具、夹具、委外和计划冲突。",
        "现场上报后自动通知、计时、升级、记录临时措施、根因、永久措施、验证和关闭。",
        "异常造成的停机、延期、返工、报废和成本进入统一分析。",
    ])


def section_qms(doc):
    add_heading(doc, "品质管理与QMS", 1)
    add_heading(doc, "全过程检验", 2)
    add_bullets(doc, [
        "IQC来料检验：材料、供应商批次、炉号、质保书、抽样、检验结果和入库/退货。",
        "首件检验：每批/每次换型首件合格后才允许批量生产。",
        "IPQC过程检验：按工序、频率、批量和风险自动生成巡检任务。",
        "末件检验：批次结束、换班或换型前检验最后加工件；不合格时反查自最近合格检验后的影响数量。",
        "FAI首件鉴定：新产品、打样、重大变更或客户要求时按气泡图记录全尺寸、材料、性能、外观和文件，并生成受控FAI报告。",
        "FQC完工检验与OQC出货检验：合格后才能入成品库或出货。",
    ])
    add_heading(doc, "检验标准与点数", 2)
    add_bullets(doc, [
        "每道工序维护检验项目、尺寸、标准值、上下限、检具、方法、频率、抽样数量和关键等级。",
        "支持计量、计数、全检、抽检、首件和破坏性检验；检验数据与工序、设备、人员、材料批次绑定。",
        "系统根据工单和工艺自动带出检验表，不由质检员自由选择旧版标准。",
        "首检、巡检/点检、末件检、FAI、完工和出货记录统一存储原始测量值、照片、附件、电子签名、报告版本和修改日志。",
    ])
    add_heading(doc, "质量异常、评审与处置", 2)
    add_bullets(doc, [
        "统一质量异常单记录来源、产品、工单、批次、当前/责任工序、设备、人员、缺陷、实测、图片、数量和影响范围。",
        "自动冻结异常批次、阻止转序/入库/出货，并检查同材料、同设备、同时间段和已出货产品。",
        "MRB处置支持返工、返修、退回上道、增加临时工序、退供应商、挑选、让步、降级和报废。",
        "跨工序返工评审必须选择“返工后返回发现工序”或“返工后从下一工序顺序重流”；选择依据、受影响特性、有效工序和失效工序均形成受控记录。",
        "让步、降级和放行必须按权限审批并记录客户确认。",
    ])
    add_heading(doc, "返工复检与质量成本", 2)
    add_bullets(doc, [
        "区分投入、一次合格、待检、不良、返工中、返工合格、返工不合格、报废、让步和入库数量。",
        "返工复检记录原不良项目、受影响关联尺寸、返工前后测量值、检具、检验人和报告。",
        "返工成本归集补料、人工、设备、刀具、工装、委外、复检、加急运输和客户补货。",
        "指标区分一次合格率和返工后最终合格率，禁止把返工合格品计入一次合格。",
    ])
    add_heading(doc, "CAPA、8D与供应商/客户协同", 2)
    add_bullets(doc, [
        "重大或重复问题启动根因分析、临时遏制、纠正、预防、责任、期限和效果验证。",
        "供应商门户接收来料/委外异常并在线回复、提交8D、补货、返工和证据。",
        "客户门户提交客诉、图片和批次，查看经授权发布的调查、返工、补货和关闭进度。",
    ])
    add_heading(doc, "检具与计量管理", 2)
    add_bullets(doc, [
        "维护检具编号、类型、精度、量程、位置、负责人、校准周期、证书、状态和适用检验项目。",
        "校准临期预警，过期/停用检具禁止使用；检验记录保留实际使用检具。",
    ])


def section_wms(doc):
    add_heading(doc, "仓储、物流与WMS", 1)
    add_heading(doc, "收货、上架和库存状态", 2)
    add_bullets(doc, [
        "供应商ASN到货后扫描送货单，核对采购/委外订单、数量、批次、箱数、质保书和预约。",
        "物料进入待检，IQC合格后推荐库位上架；不合格进入隔离区，退货或处置。",
        "库存状态区分合格、待检、冻结、不良、客户专用、供应商寄售、委外、在制和线边。",
    ])
    add_heading(doc, "仓位图和摆放标准", 2)
    add_bullets(doc, [
        "电子仓位图按仓库、库区、货架、层、位显示容量、物料、数量、批次、状态和最近盘点。",
        "每类产品维护默认库位、容器、最大数量、防混料、先进先出和标准摆放图片。",
        "颜色区分正常、接近容量、待检、冻结、不合格和空库位；点击可钻取明细。",
    ])
    add_heading(doc, "余料和多余材料管理", 2)
    add_bullets(doc, [
        "余料记录原材料批次、材质、炉号、长宽高或直径、剩余长度/重量、库位、原工单和可用状态。",
        "新订单优先匹配尺寸和材料满足的余料；下料后自动生成新余料标签。",
        "超标准损耗、报废和不可用余料必须有原因、责任和审批。",
    ])
    add_heading(doc, "生产备料与紧急配送", 2)
    add_bullets(doc, [
        "根据工单、BOM、批次和优先级生成备料、拣货、配送、补料、退料和线边仓任务。",
        "紧急缺料触发高优先级任务，记录责任人、送达时间和因缺料产生的停机。",
        "扫描物料、库位、工单和数量防止错料、错批次和错工单。",
    ])
    add_heading(doc, "盘点、准确率与权限", 2)
    add_bullets(doc, [
        "支持年度盘点、循环盘点、动态盘点和差异审批，冻结盘点范围并追溯差异来源。",
        "非仓库人员不能执行出入库、调整、盘点和报废；物理进入仓库还需门禁和现场制度配合。",
        "负库存、无单据移动和手工直接改库存原则上禁止。",
    ])
    add_heading(doc, "呆滞物料管理", 2)
    add_bullets(doc, [
        "按90/180/365天未使用、订单取消、图纸/BOM失效、客户专用无需求、超过未来需求和质量冻结识别。",
        "显示数量、金额、库龄、最后入库/领用、客户订单、未来需求和建议处置。",
        "建议处置包括优先使用、替代评审、退供应商、转售、返工利用、报废和继续保留。",
    ])


def section_finance(doc):
    add_heading(doc, "财务、业财一体化与制造成本", 1)
    add_callout(doc, "财务建设目标", "财务模块既要满足法定核算、税务、资金和内部控制，也要把报价、订单、采购、委外、库存、生产、质量、交付和售后数据转成可追溯的凭证、成本、利润与现金分析。所有财务结果必须能够钻取到原业务单据，管理口径与法定口径分层但可勾稽。")

    add_heading(doc, "财务组织、科目与核算基础", 2)
    add_bullets(doc, [
        "支持法人、账套、工厂、事业部、部门、利润中心、成本中心、工作中心和项目等多组织核算，并明确组织间交易与抵销规则。",
        "维护会计科目、辅助核算、会计期间、币种、汇率类型、结算方式、税码、银行账户、现金流量项目和凭证字。",
        "客户、供应商、员工、部门、项目、订单、产品、工单、批次、设备和工序可作为辅助核算维度，避免在科目中无限拆分。",
        "主数据新增、变更、停用必须审批；科目和核算维度启用后保留历史，不允许直接删除导致凭证失去引用。",
        "支持中国会计准则所需账簿、期间控制和凭证规则；如存在海外主体，可按主体配置本位币、会计准则、税制和报表口径。",
    ])

    add_heading(doc, "总账、凭证与期末处理", 2)
    add_bullets(doc, [
        "支持手工凭证、自动凭证、周期凭证、预提、摊销、重分类、结转、冲销、红字、外币折算和汇兑损益。",
        "凭证录入和接口生成时校验借贷平衡、期间、科目状态、辅助核算完整性、预算、现金流量和附件。",
        "凭证支持制单、复核、审核、记账、反记账、冲销和结账权限分离；高风险操作必须有审批与审计日志。",
        "提供科目余额表、明细账、总账、日记账、辅助账、序时账、试算平衡、现金流量表底稿和多维查询。",
        "期末自动执行损益结转、汇兑、折旧、成本结转、税金计提及检查；未完成前置任务时禁止关账。",
    ])

    add_heading(doc, "应收、客户信用与回款", 2)
    add_bullets(doc, [
        "销售订单、发货、签收、开票、应收、收款、退款、折让和坏账在同一链路关联，支持按订单、客户和发票核销。",
        "维护客户信用额度、信用期限、逾期容忍、预付款要求和临时信用审批；超信用或严重逾期时控制接单与发货。",
        "应收账龄按客户、业务员、区域、币种、订单和到期日展示，区分未到期、逾期区间、争议款和承诺回款日。",
        "支持预收、分期收款、质保金、尾款、票据和多笔对一笔/一笔对多笔核销，保留差异、短款和手续费。",
        "形成催款任务、客户对账单、回款计划、逾期预警和坏账准备建议，并记录催收过程与责任人。",
    ])

    add_heading(doc, "应付、供应商发票与付款", 2)
    add_bullets(doc, [
        "采购和委外按订单、收货/回厂、合格数量、退货、扣款、对账、发票和付款形成完整三单或四单匹配。",
        "采购对账单与委外对账单经供应商在线确认后生成暂估或正式应付，差异必须进入审批和原因记录。",
        "支持预付、到货款、进度款、质保金、尾款、代扣、索赔、返利、票据和多币种付款。",
        "付款申请依据到期日、资金计划、供应商等级、停供风险、折扣和争议状态排序，禁止重复付款和超额付款。",
        "供应商发票执行订单、收货、税率、金额、抬头、发票号码和查重校验；红冲、退票和重开发票与原业务关联。",
    ])

    add_heading(doc, "现金、银行与资金计划", 2)
    add_bullets(doc, [
        "维护银行账户、现金账户、网银权限、账户余额、受限资金、票据、授信、贷款、保函和到期计划。",
        "银行流水可通过银企直联或文件导入，按金额、日期、对方和摘要智能匹配收付款单，未匹配项形成待办。",
        "滚动资金预测汇总销售回款、采购付款、工资、税费、设备投资、贷款和其他计划，形成日、周、月现金缺口。",
        "付款批次支持制单、复核、审批、出纳支付和结果回写分离，付款文件、回单和失败原因自动归档。",
        "提供账户余额、可用资金、未来收支、资金缺口、融资需求、逾期债务和大额资金变动预警。",
    ])

    add_heading(doc, "费用、报销与员工往来", 2)
    add_bullets(doc, [
        "支持差旅、招待、交通、采购、质量索赔、加急物流、样品、工装和其他费用申请、借款、报销与还款。",
        "费用单关联预算、部门、项目、客户、订单、工单、质量异常和成本对象，附件包括发票、合同、行程和审批证据。",
        "校验费用标准、超标原因、重复发票、同一发票多次报销、借款未清和个人收款账户。",
        "公司卡、员工借款、备用金和费用报销自动核销；长期未清借款和缺失附件触发预警。",
        "费用审核通过后自动生成应付或付款及会计凭证，并进入部门、订单、项目和客户获利分析。",
    ])

    add_heading(doc, "固定资产、低值易耗品与租赁", 2)
    add_bullets(doc, [
        "资产从采购、验收、转固、编码、位置、责任人、使用状态到调拨、维修、盘点和处置全过程管理。",
        "支持资产类别、原值、残值、折旧方法、使用年限、累计折旧、减值、改良和多账簿核算。",
        "设备台账与固定资产卡片建立映射，维修成本、停机、产能和折旧可共同进入设备收益分析。",
        "低值易耗品、工具和办公资产支持领用、归还、摊销和盘点；租赁资产支持合同、付款和到期提醒。",
        "资产新增、调拨、报废、出售和盘亏必须审批，并自动生成资产与总账凭证。",
    ])

    add_heading(doc, "存货核算与业财对账", 2)
    add_bullets(doc, [
        "存货按物料、仓库、库位、批次、所有权和库存状态核算，支持移动平均、月末一次加权或企业选定的合法计价方法。",
        "采购入库、暂估、发票、退货、调拨、领退料、完工入库、销售出库、盘盈盘亏和报废自动形成存货核算记录。",
        "区分合格、待检、冻结、不良、客供、寄售、委外、在制和线边库存的数量与价值，明确是否纳入可用和财务资产。",
        "库存数量账、存货核算账和总账自动勾稽；数量为零金额不为零、负库存、负金额和跨期单据形成异常清单。",
        "期末暂估差异、采购价格差异、材料价差、盘点差异和报废损失按规则结转并可追溯原单。",
    ])

    add_heading(doc, "税务、发票与外贸财务", 2)
    add_bullets(doc, [
        "维护销项和进项税码、税率、生效日期、含税/未税口径、可抵扣规则和税务组织，业务单据自动带出并允许受控调整。",
        "销售开票关联订单、发货、签收和应收；采购发票关联订单、收货、对账和应付，支持蓝票、红票、部分开票与差额。",
        "建立发票号码、代码、抬头、税号、金额和影像的查重及异常校验，发票状态与应收应付同步。",
        "外贸支持本位币与交易币、即期/记账/结算汇率、汇兑损益、出口发票、收汇和退税资料关联。",
        "Commercial Invoice、Packing List、报关数据、销售合同、物流、出口收款和会计凭证按同一出货批次追溯。",
    ])

    add_heading(doc, "香港代生产客户正式订单70%价格财务控制", 2)
    add_bullets(doc, [
        "财务审核时先校验客户主数据勾选状态与订单类型：只有“已勾选客户＋正式业务订单”才能使用70%计算后价格。",
        "系统并列保存业务输入的原始价格、固定70%系数和计算后正式订单价格；原始价格用于追溯，计算后价格作为该订单后续普通应收、成本毛利和经营分析的业务价格。",
        "样品订单、模具订单和未勾选客户均使用原输入价格，不得因客户名称、香港地址或人工备注自动触发70%转换。",
        "除价格计算外，不建立30%留存、公司间分成或特殊结算模块；应收、开票、收款、汇兑和总账继续按普通订单既有规则处理。",
        "订单全部材料、人工、机时、制造费用、委外、返工、报废、检验、包装、物流和出口费用继续完整归集，并以计算后正式订单价格分析毛利。",
        "客户、订单类型、原始单价、数量或币种变化时自动重算；已审核或已开票订单必须发起普通订单变更、重新审核并保留历史价格版本。",
    ])

    add_heading(doc, "预算、资金占用与费用控制", 2)
    add_bullets(doc, [
        "预算可按年度/月度、组织、部门、项目、科目、客户、产品和设备编制，支持自上而下分解与自下而上汇总。",
        "采购申请、费用申请、付款、资本支出和合同占用预算，执行、释放、追加、调剂和超预算审批均保留记录。",
        "预算执行看板显示预算、占用、实际、可用、预测和差异，并可钻取订单、费用和凭证。",
        "资金占用分析覆盖库存、在制、应收、预付、固定资产和呆滞物料，识别占用金额、天数和责任部门。",
        "预算口径、版本和调整原因受控，未经审批不得通过修改历史预算掩盖超支。",
    ])

    add_heading(doc, "业务单据自动凭证与自动对账", 2)
    add_bullets(doc, [
        "采购入库、暂估、发票、退货、销售出库、销售发票、收付款、生产领退料、完工入库、委外、报废、费用和资产按规则自动生成凭证。",
        "自动凭证模板维护来源事件、借贷科目、辅助核算、金额来源、税额、摘要、现金流量、组织和生效版本。",
        "凭证可双向追溯原业务单据；业务数量金额与库存、应收、应付、资产、成本和总账之间每日或月末自动对账。",
        "接口失败、科目缺失、辅助核算不完整、金额不平和期间关闭时进入异常队列，不得静默丢失。",
        "冲销或更正必须从源业务单据发起并形成反向凭证，禁止只改总账而保留错误业务状态。",
    ])

    add_heading(doc, "制造成本对象与归集规则", 2)
    add_bullets(doc, [
        "成本对象支持产品、客户、销售订单、生产工单、工序、项目、批次、设备、工作中心和部门。",
        "直接材料依据实际领退料、材料批次和计价成本归集；超额领料、替代料、余料退库和材料损耗单独识别。",
        "直接人工依据审核后的报工工时、人员类别和工价归集；机器成本依据设备运行工时、标准费率、折旧、能耗和维修分摊。",
        "委外成本依据委外订单、发料、合格回厂、损耗、返工、扣款和对账归集；刀具、工装、检验、包装和物流按规则分配。",
        "制造费用按部门、工作中心、设备或工时动因分配，分配基准、费率和生效期间必须有版本。",
    ])

    add_heading(doc, "工序级制造成本核算：每一道工艺独立计价", 2)
    add_callout(doc, "强制颗粒度", "成本必须核算到订单→工单→批次/序列号→路线版本→工序→工序执行实例。切料、CNC、车床、打磨、清洗、振磨、委外、镭雕、丝印、组装、检验、全检和包装不得合并成一个“加工费”；返工或重复执行必须生成第二次工序成本。", fill="FFF4E5", color=GOLD)
    add_heading(doc, "工序成本核算原则", 3)
    add_numbered(doc, PROCESS_COST_PRINCIPLES, font_size=9.0, space_after=3, line_spacing=1.12)
    add_heading(doc, "工序成本卡字段", 3)
    add_table(doc, ["数据分组", "字段与统一口径"], PROCESS_COST_FIELDS, [2050, 7310], font_size=8.25)
    add_heading(doc, "每道工序必须归集的成本", 3)
    add_table(doc, ["成本组", "必须归集的内容"], PROCESS_COST_COMPONENTS, [2050, 7310], font_size=8.25)
    add_heading(doc, "工序成本计算与累计流转", 3)
    add_table(doc, ["指标", "公式/计算", "系统控制"], PROCESS_COST_FORMULAS, [1850, 3900, 3610], font_size=7.95)
    add_heading(doc, "各工艺成本采集矩阵", 3)
    add_table(doc, ["工艺", "独立成本内容", "主要数据来源"], PROCESS_COST_MATRIX, [1650, 4300, 3410], font_size=7.8)
    add_heading(doc, "工序成本形成、结算与重算流程", 3)
    add_numbered(doc, PROCESS_COST_WORKFLOW, font_size=8.9, space_after=3, line_spacing=1.1)
    add_heading(doc, "工序成本审核权限", 3)
    add_numbered(doc, PROCESS_COST_APPROVALS, font_size=8.9, space_after=3, line_spacing=1.1)
    add_heading(doc, "工序成本专项预警", 3)
    add_bullets(doc, PROCESS_COST_ALERTS)
    add_heading(doc, "工序成本专项报表", 3)
    add_table(doc, ["报表", "系统生成内容"], PROCESS_COST_REPORTS, [2500, 6860], font_size=8.2)

    add_heading(doc, "标准成本、实际成本与在制分配", 2)
    add_bullets(doc, [
        "标准成本来自有效BOM、工艺标准工时、材料标准价、人工费率、设备费率、委外价和费用分摊规则。",
        "材料标准价可引用实时金属价格中心的受控日结/周期快照，并叠加企业升贴水与加工物流；价格更新先生成模拟差异，审批后才形成新标准成本版本。",
        "实际成本按工单归集材料、人工、机时、委外和费用，支持分批完工、联副产品、在制约当量和跨月工单。",
        "期末区分完工、在制、报废、返工和待判数量，按企业批准的方法分配成本并保留计算底稿。",
        "分析材料价格差异、材料用量差异、人工工资率差异、人工效率差异、设备效率差异、委外价格差异和制造费用差异。",
        "按单件、批次、工单、订单和客户比较报价成本、标准成本、实际成本、销售价格和实际毛利。",
    ])

    add_heading(doc, "返工、报废与质量损失成本", 2)
    add_bullets(doc, [
        "内部返工和客户退回返工分别归集补料、人工、机时、刀具、工装、委外、复检、物流、补货和延期损失。",
        "跨工序路线A只归集第二工序返工及复检新增成本；路线B还归集第三、第四工序重新执行的人工、机时、材料和检验成本。",
        "原第三、第四工序被路线B替代的历史成本不删除，应按会计政策转入质量损失、在制调整或责任分析，避免重复计入合格品成本。",
        "报废记录发生节点、累计在制成本、可回收材料价值、处置收入、责任工序、责任部门和审批。",
        "质量成本分为预防、鉴定、内部失败和外部失败，可按产品、客户、供应商、部门、缺陷和责任原因分析。",
    ])

    add_heading(doc, "部门核算、利润中心与经营分析", 2)
    add_bullets(doc, [
        "按部门、工作中心、设备、项目或利润中心归集收入、材料、人工、机时、费用、返工、报废和质量损失。",
        "内部结算可按标准费率、实际服务量或协议价在部门间分配，规则必须透明、有版本并与法定账分层。",
        "形成客户、产品、订单、项目、设备和部门损益，识别低毛利订单、亏损客户、低效设备和高损耗工序。",
        "贡献毛利、完全成本毛利、EBITDA类管理指标必须有统一定义，禁止各部门使用不同Excel口径。",
        "管理报表与财务报表可以采用不同展示口径，但必须通过调整项目和明细钻取完成勾稽。",
    ])

    add_heading(doc, "月结、财务报表与审计档案", 2)
    add_bullets(doc, [
        "月结任务清单覆盖库存关账、采购暂估、发票、应收应付、成本计算、折旧、费用、税务、汇兑、对账和总账结转。",
        "每项任务维护责任人、前置条件、计划完成时间、实际时间、异常和审批；临近和超过时限进入统一预警中心。",
        "系统生成资产负债表、利润表、现金流量表、所有者权益变动及管理报表，并支持组织、期间和版本对比。",
        "提供业务到财务、明细到总账、管理报表到法定报表的勾稽关系和差异解释。",
        "凭证、合同、发票、回单、对账单、成本底稿、审批和报表快照形成电子会计档案，保留期限、访问和下载可审计。",
    ])

    add_heading(doc, "财务与老板AI问答：总体架构", 2)
    add_bullets(doc, [
        "在财务工作台和老板经营驾驶舱分别提供自然语言问答入口，支持连续追问、条件修改、结果收藏和导出分析报告。",
        "AI只能通过经审批的指标语义层、只读查询服务和受控知识库访问ERP、MES、QMS、WMS、SRM、客户门户及财务数据，不直接连接生产数据库执行任意写入。",
        "语义层统一收入、成本、毛利、现金、应收、应付、库存、在制、准交、良率和OEE等指标的定义、维度、时间口径和负责人。",
        "知识库纳入会计政策、预算制度、费用标准、合同模板、指标字典和已批准SOP；文件内容仅作为资料，不得被当作改变系统权限或执行操作的指令。",
        "每次回答记录提问人、角色、模型/版本、时间、数据截止时间、查询条件、引用数据、生成结果、导出和反馈，满足审计与问题复盘。",
    ])

    add_callout(doc, "模型选型", "AI问答模型采用DeepSeek V4 Pro，通过DeepSeek官方API接入（内部运维调试可用DeepSeek TUI，业务用户仅经系统内问答界面）。模型只承担自然语言理解、语义层查询生成和答案组织，数值一律来自语义层查询结果；API密钥统一托管，请求前对工资、银行账号等敏感字段脱敏；每次调用记录模型名称、版本、token用量和费用；模型或提示模板升级前用固定财务问题集回归；模型不可用时自动降级为标准报表和驾驶舱，不阻塞业务。", fill="EAF2FB", color=DARK_BLUE)
    add_heading(doc, "财务人员AI分析场景", 2)
    add_bullets(doc, [
        "应收分析：询问本周应收、逾期客户、金额、账龄、承诺回款和催收优先级，并下钻到订单、发票、收款和业务员。",
        "应付与资金：询问未来7/30/90天付款需求、可用资金、预计回款、资金缺口和建议付款顺序，展示假设而不是自动支付。",
        "成本差异：询问某订单为何实际成本高于报价或标准，拆解材料价格/用量、人工、机时、委外、返工、报废和费用差异。",
        "工序成本：询问该订单每道工艺分别花了多少钱、哪一道成本占比最高、CNC/打磨/委外/镭雕/丝印/检验/包装为何超标准，并下钻报工、设备、材料和委外源单。",
        "毛利分析：询问哪些客户、产品、订单或业务员毛利下降，区分价格、材料、效率、质量损失、汇率和产品结构因素。",
        "月结检查：询问还有哪些未结任务、未过账单据、暂估、负库存、对账差异和成本异常，并生成责任人清单。",
        "发票与税务：询问未开票发货、发票与订单/收货差异、重复发票、进项待认证和出口收汇/退税资料缺口。",
        "预算与费用：询问超预算部门、异常增长科目、重复或异常报销、预算占用和未来现金影响。",
        "材料行情：询问铜、铝、钢材或其他金属价格变化对未报价询价、已报价未成交订单、在途采购、库存、在制、标准成本和毛利的影响。",
        "香港客户价格：询问哪些正式订单因客户勾选应用了70%价格、原始价格与计算后价格是多少、是否存在漏算/重复折算、实际成本和毛利如何，并引用客户主数据、订单和价格版本。",
        "所有AI结论均可一键生成带数据截止时间、口径、图表、明细和引用单据的财务分析报告，但凭证、付款、关账和报税仍须按权限人工确认。",
    ])

    add_heading(doc, "老板AI经营问答场景", 2)
    add_bullets(doc, [
        "经营总览：询问今天、本周、本月和年度的订单、销售、回款、毛利、现金、库存、在制、交付、质量和设备情况，并与目标、预算、上期和去年同期比较。",
        "原因追问：从“本月毛利为什么下降”继续追问到客户、产品、订单、材料、工序、供应商、返工和报废，形成可验证的因果线索。",
        "工艺盈利：询问不同产品/客户的各工艺成本、单位成本和标准实际差异，识别需要调价、改善效率、换机或调整委外策略的具体工序。",
        "交付风险：询问未来两周哪些订单可能延期，展示缺料、供应商、工序、设备、质量和审批原因及预计影响金额。",
        "现金与风险：询问可用现金、未来收支、逾期应收、集中到期应付、客户信用、库存占用和潜在资金缺口。",
        "客户与供应商：询问客户贡献、订单趋势、毛利与回款质量，以及供应商价格、准交、质量、风险和替代来源。",
        "情景模拟：在明确假设下分析材料涨价、汇率变化、插单、设备停机、回款延迟或采购提前对毛利、交期和现金的影响；模拟结果必须与正式账表分开标识。",
        "材料决策：询问近期金属行情、供应商报价与市场偏离、重点订单价格风险及采购/报价窗口，AI必须引用许可行情快照而不是自行猜测市场价格。",
        "香港客户经营：询问被勾选客户的正式订单原始价格、70%后业务价格、订单额、成本毛利、回款、交付和客户/产品贡献；样品与模具订单单独显示且证明未应用70%。",
        "老板看到的汇总与钻取范围由权限决定；工资明细、个人银行信息、供应商底价和未授权客户数据不得因AI问答而越权暴露。",
    ])

    add_heading(doc, "AI回答标准、权限与风险控制", 2)
    add_bullets(doc, [
        "每个答案必须显示数据截止时间、组织/期间/币种/含税口径、指标定义、筛选条件、计算逻辑和来源系统。",
        "金额和关键指标必须提供可点击引用，能够钻取到报表、凭证或业务单据；无法取得数据时明确说明缺口，不得编造。",
        "AI应区分已验证事实、基于数据的推断、预测和建议；预测需显示模型、假设、区间或不确定性。",
        "复用ERP的组织、行级、字段级和敏感数据权限；管理员也不能通过提示词绕过薪资、银行、成本、供应商价格和客户隔离。",
        "AI默认只读，不能自动审核、记账、付款、关账、修改主数据或对外发送文件；如未来开放写回，必须生成草稿、二次确认并走原审批流。",
        "建立敏感信息脱敏、传输加密、日志留存、模型供应商数据不训练承诺、知识库版本、内容安全和人工复核机制。",
        "验收时使用固定问题集与人工核算结果逐项对比，验证数字一致、口径透明、权限隔离、追问稳定、引用有效和响应时间。",
    ])


REPORT_SPECS = [
    ("客户报价单", "报价审批通过或授权人员一键生成", "客户、报价版本、图纸、数量阶梯、价格、税率、币种、交期、付款、包装运输和条款", "中文/英文PDF；客户版隐藏内部成本；保留发送和版本记录"),
    ("内部成本分析表", "工程核价完成或报价版本变更", "材料尺寸重量、单价、利用率、工序机时、人工、刀具、工装、委外、检验、包装、费用和毛利", "PDF/Excel；订单完成后可与标准/实际成本对比"),
    ("模具/样品/正式订单收费与成本报告", "订单审批、成本发生、发货、开票、回款、抵扣、返还或成本分摊后实时更新", "订单类型、收费方式、客户应付、免费/减免、预计/实际成本、承担部门、模具所有权、样品转化、正式订单关联、已/未回收成本和应收回款", "按三类订单分别统计；免费订单显示审批额度和成本；正式订单显示收费、应收和回款，禁止混入免费口径"),
    ("香港客户正式订单价格计算与审计表", "客户价格勾选变更、正式订单提交、客户/订单类型/原价/数量/币种变更或审核时", "客户勾选快照、订单类型、原始单价金额、适用系数、70%后单价金额、版本、审核、下游引用及异常", "PDF/Excel；证明仅已勾选客户的正式订单应用70%，样品/模具和其他客户保持原价；不另建分成或特殊结算表"),
    ("外贸发票与报关资料", "发货审核、箱数和重量确认后", "客户英文资料、HS编码、品名、原产国、币种、贸易条款、数量、金额、毛净重、箱号和运输", "PI、Commercial Invoice、Packing List、销售合同及报关数据包"),
    ("订单全生命周期时间轴", "业务订单提交时建立；每次领取、审核、开工、完工、送检、交接、回料和包装后实时更新", "T0/T1、节点基线/最新计划/实际时间、首次响应、排队、处理、暂停、总历时、超期、责任人、退回轮次、阻塞和预计完成", "在线看板、PDF/Excel快照；逐单钻取全部事件、单据、审批和异常"),
    ("审核与节点时效分析", "实时看板；日/周/月定时快照", "业务经理审核、财务业务订单审核、订单评审、MRP、采购、财务采购单审核、收料、IQC、各工序、委外和包装的平均/中位/P90/P95、按时率、超期率和积压", "按部门、岗位、人员、客户、订单、节点、供应商和紧急度下钻，形成瓶颈Pareto"),
    ("工序数量交接差异表", "每次工序交接确认或出现短少/多出后实时更新", "订单、工单、批次/序列号、交出/接收工序、申报数、实点数、实转数、差异数、双方人员、时间、原因、责任和关闭", "差异未关闭持续预警；与质量NCR分离；显示对在制、齐套和交期的影响"),
    ("订单进度表", "实时滚动；每日定时快照", "订单数量、在制、入库、出货、当前工序、工时权重进度、齐套、质量、预计完成和交付风险", "在线看板、PDF/Excel快照；可钻取工单和异常"),
    ("工序进度表", "扫码开工/暂停/完工/返工后实时更新", "工序投入、一次合格、返工、报废、剩余、计划/实际时间、设备、人员、超期和等待原因", "按车铣磨、热处理、表面处理、装配、检验、返工等视角输出"),
    ("采购进度表", "采购申请、订单、确认、ASN、到货、IQC和入库变化后", "需求、库存、缺口、订单、供应商、承诺/预计交期、到货、合格、不良、未交和影响工单", "在线滚动、紧急清单、超期清单和每日快照"),
    ("PMC成品备料与齐套表", "订单、BOM、库存、报工、报废、冻结、采购/委外变化后", "毛需求、损耗、合格库存、分配、待检/冻结、在制、在途、缺口、建议采购/生产、关键料和最晚齐套", "订单/成品/组件多层展开，显示数量/品种/关键料齐套"),
    ("供应商进度表", "供应商确认、填报进度、ASN、到货和质量变化后", "采购/委外订单、承诺交期、完成、已交、未交、预计交期、超期、质量和影响工单", "供应商、采购员、物料和紧急程度多视角"),
    ("供应商报价与比价表", "RFQ截止、每轮报价和定标时", "阶梯价、税率、运费、工装、交期、产能、付款、技术偏离、质量与交付绩效", "供应商报价PDF、综合比价、技术偏离和定标审批单"),
    ("采购对账单", "对账周期或供应商发起", "采购订单、合格入库、退货、价格、税率、发票、已/未开票和差异", "供应商在线确认；差异审批；确认后形成应付"),
    ("委外对账单", "对账周期或委外订单完成", "委外订单、发料、收料、合格、返工、报废、损耗扣款、加工费、运输和税", "供应商在线确认；确认后形成应付"),
    ("仓位图", "库存、上架、移库、盘点和冻结后实时更新", "仓库/库区/货架/层/位、容量、物料、批次、数量、状态和盘点时间", "可交互电子图；颜色区分空位、待检、冻结、不良和满载"),
    ("首检/巡检/末件检验报告", "对应检验任务完成、复核或版本发布后", "订单、产品、图纸、工单、工序、气泡号、标准/实测、检具、设备、人员、程序、时间、图片和结论", "系统存储原始测量值和签名；PDF/Excel按类型、批次和版本生成"),
    ("FAI报告", "新产品、打样、重大变更或客户要求的FAI审核通过后", "气泡图、全尺寸、材料、性能、外观、过程、检具、实测、偏差、批准和附件", "客户/内部模板PDF；报告版本与图纸、样品工单和变更绑定"),
    ("后工序履历与参数报告", "清洗、丝印、镭雕、打磨或振磨扫码完工后", "工序参数、材料/介质批次、程序/图稿、设备、人员、时间、投入、合格、不良、返工和检验", "按批次/单件展示完整后工序链和版本，支持异常钻取"),
    ("出货前全检与包装报告", "100%全检、OQC、标签和包装任务全部完成后", "逐件/逐批检验、外观/关键尺寸、标识、数量、箱号、标签、重量、包装规范、照片和文件齐套", "客户/内部PDF；未完成或不合格时禁止生成可发货状态"),
    ("组装工单与组件谱系报告", "组装领料、装配、过程检查、功能测试和完工后", "组装BOM、组件批次/序列号、装配顺序、扭矩/参数、辅料、人员、工位、测试、返工和成本", "成品到组件正反追溯；组装记录、测试报告和完工清单"),
    ("品质异常分析表", "质量异常创建、处置、返工、CAPA和关闭后", "产品、工单、工序、设备、人员、材料批次、缺陷、数量、图片、原因、责任、时长、成本和重复发生", "Pareto、趋势、责任/设备/供应商排行和明细钻取"),
    ("成本核算分析", "工单完工、月结或成本重算后", "报价、标准、实际材料/人工/机时/委外/返工/报废/费用、差异和毛利", "产品、订单、工单、工序、客户、设备、部门多维分析"),
    ("工序成本明细卡与成本瀑布", "每道工序开工建立；报工、领退料、设备采集、检验、交接、委外、费用分配或重算后实时更新", "订单、工单、批次、路线版本、工序执行实例、转入累计成本、本工序材料/人工/机器/刀具工装/委外/检验/物流/制造费用/异常/回收、转出与累计成本", "逐工序成本卡及从切料到包装的成本瀑布；订单总成本必须等于全部有效工序实例汇总并可下钻源单"),
    ("工序标准实际成本差异分析", "工序结算、工单完工、月结或受控重算后", "工序标准/预计/实际暂估/实际结算成本、材料价差/量差、人工工资率/效率、设备效率、委外价格、良率、返工报废和费用分配", "按产品、订单、批次、工序、设备、人员、班次、供应商和期间比较；显示费率版本与数据截止时间"),
    ("在制品工序成本报告", "工序报工、暂停、跨月、拆批/合批、质量冻结或结算变化后", "在制所在工序、转入累计成本、本工序已发生未完工成本、等价产量、预计完工成本、停留时长和阻塞", "数量与金额同时展示；按订单/工序/批次/工作中心下钻，支持月末在制分配底稿"),
    ("工序成本完整性与月结检查", "实时检查及工单结算/月结前强制运行", "未配置成本模板/费率、未报工、零负成本、数量不平、成本链断裂、重复分配、未结委外、未分制造费用、待重算和异常调整", "形成责任人待办和阻断清单；问题全部关闭后才能锁定工序成本、结算工单或关账"),
    ("停机原因分析", "设备自动采集或人工登记停机后", "设备、开始/结束、时长、原因、工单、人员、责任、影响交期和维修", "停机原因Pareto、设备排行、趋势和计划/非计划比例"),
    ("报废分析", "报废申请、审批和成本结转后", "产品、工单、工序、批次、设备、人员、原因、数量、累计成本、可回收价值和责任", "数量、金额、原因、工序和责任趋势"),
    ("良率分析", "每次检验、报工和返工复检后", "一次合格、返工后最终合格、工序/产品/设备/人员/供应商良率", "明确排除返工合格对一次合格率的虚增"),
    ("材料损耗分析", "领退料、余料、报废和工单完工后", "标准用量、实际领料、退料、正常/超额损耗、余料、废料、回收和责任工序", "材料利用率、损耗金额、订单/产品/工序排行"),
    ("仓库呆滞物料分析", "每日/每周定时及订单/BOM失效后", "物料、批次、库位、数量、金额、库龄、最后收发、客户订单、未来需求和处置建议", "90/180/365天、订单取消、版本失效、客户专用、余料不可匹配等口径"),
    ("金属材料价格与影响分析", "行情更新、供应商报价、采购成交、汇率变化或授权人员生成时", "市场基准、供应商报价、采购成交、汇率、升贴水、落地价、报价/订单/库存/在制/采购承诺和毛利影响", "采购/财务/业务权限视图；显示行情源、许可、实时性、快照和换算公式"),
    ("考勤日报/月报", "每日计算、月度封存或授权人员生成时", "应出勤、实出勤、正常/加班工时、请假、迟到、早退、缺卡、异常、班次、部门、用工类型和劳务单位", "个人/班组/部门/劳务多视图；封存版本、审批和调整记录"),
    ("考勤与MES工时差异报告", "考勤与报工每日对比或月度结算前", "打卡、有效报工、计件、跨部门借调、无报工、无打卡、超长工时、责任人和处理状态", "异常待办、原因、审批和关闭；不得自动改工资或报工"),
    ("劳务对账与结算单", "考勤封存、部门确认及劳务结算周期到达后", "劳务单位、人员、出勤、正常/加班小时、班次、岗位、计件、费率、管理费、补贴、扣款、税费、发票和差异", "行政/部门/财务审批；劳务确认；转应付并保留合同费率版本"),
    ("跨工序返工路线与成本分析", "返工路线审批、每道返工工序报工/复检及返工关闭后", "原路线、返工起点、路线A/路线B、子批次、失效/重做工序、复检、工时、材料、机时、交期和质量损失", "完整路线图、数量状态和新旧履历；可比较直接返回与顺序重流成本"),
    ("返工子订单成本明细表", "NCR/MRB批准返工时建立返工子订单；每次领料、报工、送检、委外、回料、包装、索赔或关闭后实时更新", "原订单、返工子订单、责任版本、受影响数量、返工路线、重复工序、直接材料、人工、机时、制造费用、检验、包装、物流、委外应付、延误与产能机会损失、赔偿/扣款/保险回收", "按返工批次分别显示原订单成本、返工新增毛成本、已确认回收和企业净损失；可逐笔钻取报工、领料、委外订单、应付、NCR和责任认定"),
    ("供应商质量损失与索赔报告", "供应商责任初判、终判、无偿返工、补货、扣款、索赔、对账或结案后实时更新", "供应商、采购/委外订单、NCR、责任比例、市场等价返工费、实际追加应付、退运物流、复检、重做后续工序、重新包装、报废、停线、延期、客户索赔、已索赔/已抵扣/已收回", "同时展示供应商追加应付、供应商造成的企业毛损失、可追偿、已回收和净损失；供应商责任无偿返工时追加应付必须为零，但质量损失不能记为零"),
    ("AI问答分析报告", "财务或老板在AI问答中确认保存/导出时", "原问题、追问、数据截止时间、口径、筛选、指标、图表、引用单据、事实/推断/预测标识和生成者", "受权限控制的PDF/Excel；保留模型版本、审计日志和引用链接"),
]


def section_documents(doc):
    add_heading(doc, "系统单据与报表中心", 1)
    add_callout(doc, "硬性原则", "Excel只能作为系统导出格式，不能作为进度、库存、对账或成本的事实来源。正式文件如需修改，必须修改源业务数据、重新审批并生成新版本，不得直接修改PDF。")
    add_heading(doc, "模板、生成与版本", 2)
    add_bullets(doc, [
        "模板管理公司抬头、Logo、语言、客户专版、币种、税率、页眉页脚、电子签章、签字人、编号和有效期。",
        "支持业务事件触发、定时生成、实时看板和授权用户一键生成四种模式。",
        "每份正式文件保存来源单据、数据截止时间、生成时间、模板、版本、审批、发送、下载、打印和失效状态。",
        "V1修改为V2时保留V1并标记失效；历史文件数据快照不能被当前业务数据反向改变。",
        "所有运营报表显示最后更新时间、数据截止时间、筛选条件和指标口径，并可钻取原始业务明细。",
    ])
    add_heading(doc, "系统生成文件目录", 2)
    add_table(doc, ["文件/报表", "触发方式", "主要数据", "输出与控制"], REPORT_SPECS, [1500, 1800, 3300, 2760], font_size=7.8)

    for name, trigger, data, output in REPORT_SPECS:
        add_heading(doc, name, 2)
        add_para(doc, f"生成时机：{trigger}", bold_prefix="生成时机：")
        add_para(doc, f"数据范围：{data}", bold_prefix="数据范围：")
        add_para(doc, f"输出要求：{output}", bold_prefix="输出要求：")
        if name in ("订单全生命周期时间轴", "审核与节点时效分析", "工序数量交接差异表", "订单进度表", "工序进度表", "采购进度表", "PMC成品备料与齐套表", "供应商进度表"):
            add_bullets(doc, ["该报表必须随业务扫码、到货、检验、入库、报废、冻结或交期变化自动滚动，禁止由部门人工填写完成比例。", "临期、超期、缺料或质量冻结数据须与统一预警中心联动，并明确责任人和预计恢复时间。"])


def section_alerts(doc):
    add_heading(doc, "统一预警、待办与升级中心", 1)
    add_heading(doc, "预警级别与时间口径", 2)
    add_table(
        doc,
        ["级别", "含义", "处理要求"],
        [
            ("蓝色", "正常待办", "按计划执行并可查询剩余时间"),
            ("黄色", "临期", "提醒责任人，确认是否存在阻塞"),
            ("橙色", "即将超期或高风险", "要求更新原因和预计完成时间"),
            ("红色", "已超期", "通知责任人和直属主管并进入超期看板"),
            ("深红色", "严重超期", "升级部门负责人/厂长/管理层并持续催办"),
        ],
        [1200, 3000, 5160],
        font_size=9,
    )
    add_bullets(doc, [
        "规则支持自然日、工作日、班次和小时，引用企业工作日历、节假日、班次和订单紧急等级。",
        "每条预警包含对象、触发时间、截止时间、超期时长、责任人、升级人、影响订单、原因、措施和关闭验证。",
        "责任人不能删除预警，只能确认、处理、申请关闭或按权限转交。",
    ])
    add_heading(doc, "工序和交付预警", 2)
    add_bullets(doc, [
        "计划开始时间到达仍未开工、长时间无报工、实际工时超标准、进度落后、计划结束超期。",
        "上道延误、缺料、设备、质量、工装、人员和委外异常造成的预计延期。",
        "图纸、工艺、程序、材料、检验、包装、物流未齐套或预计无法满足客户交期。",
        "预警展示订单、工单、工序、超期、责任、原因、后续影响和预计恢复。",
    ])
    add_heading(doc, "审批超期预警", 2)
    add_bullets(doc, [
        "覆盖报价核价、报价审批、合同评审、订单变更、图纸/BOM/工艺、工程变更、采购/委外、质量评审、返工路线、让步、报废、8D、付款和费用审批。",
        "每类审批配置SLA，临期提醒、超期升级、严重超期再次升级；请假时支持授权代理。",
        "记录送审、首次查看、审批、超期、催办、转交和驳回时间，用于审批效率分析。",
    ])
    add_heading(doc, "专项预警", 2)
    add_table(
        doc,
        ["领域", "典型预警"],
        [
            ("采购/物料", "安全库存、MRP缺料、订单未确认、到货超期、IQC等待、关键料晚于开工、余料/呆滞"),
            ("金属行情", "价格涨跌超阈值、供应商报价偏离、报价快照过期、成本/毛利风险、行情延迟中断和许可到期"),
            ("委外", "未收料、未开工、进度落后、临期、超期、回厂不良和对账未确认"),
            ("质量", "首件/巡检/末件/FAI等待或漏检、异常未判定、返工/复检超期、CAPA未关闭、重复缺陷、检具校准"),
            ("后工序/组装", "清洗/丝印/镭雕/打磨/振磨参数缺失、组装缺料/测试不合格、全检包装未完成或超期"),
            ("设备", "点检/保养到期、报修未响应、维修超期、连续停机、节拍或良率异常"),
            ("门户", "供应商未报价/确认、送货未到、客户待确认交期/图纸、外部质量回复超期"),
            ("财务", "应收逾期、发票差异、对账未确认、成本结算异常和月结任务超期"),
            ("工序成本", "工序成本模板/费率缺失、漏报工、零负成本、数量成本不平、成本链断裂、委外未暂估、返工重复成本遗漏"),
            ("香港客户价格", "正式订单70%漏算/重复折算、样品或模具误算、客户或类型变更未重算、价格版本不一致、计算后毛利过低"),
            ("行政", "缺卡、未审批加班/请假、考勤与MES工时差异、劳务合同/费率到期、对账或结算超期"),
        ],
        [1600, 7760],
        font_size=9,
    )
    add_heading(doc, "预警处理闭环", 2)
    add_numbered(doc, [
        "系统按事件或定时规则触发预警并通知责任人。",
        "责任人确认，选择标准原因或补充说明，填写临时措施和预计恢复时间。",
        "超过SLA未处理时按岗位、部门和订单等级自动升级。",
        "业务完成后提交关闭，必要时由品质、PMC或主管验证。",
        "关闭数据进入超期、根因、责任、处理时长和重复问题分析。",
    ])


def section_portal_security(doc):
    add_heading(doc, "客户与供应商门户、安全和权限", 1)
    add_heading(doc, "外部账号与数据隔离", 2)
    add_bullets(doc, [
        "每家供应商/客户可建立管理员、业务、计划、送货、质量和财务等多个账号。",
        "采用组织、业务对象和行级权限：供应商只能看自己的询价、报价、订单、质量和对账；客户只能看自己的项目和文件。",
        "合作终止、人员离职、资质过期时账号自动或人工停用；支持密码策略、多因素认证和异常登录提醒。",
        "图纸和正式文件支持水印、有效期、下载权限、下载日志和失效版本提示。",
    ])
    add_heading(doc, "门户消息与通知", 2)
    add_bullets(doc, [
        "供应商：新询价、报价截止、议价、订单变更、交期临近、送货、质量、对账、发票和付款通知。",
        "客户：新报价、订单确认、图纸/样品确认、交货计划变更、发货、质量文件、客诉和返工进度。",
        "被勾选客户仍使用普通客户门户查看自己的订单、交付与正式文件；70%价格规则不增加独立门户、分成确认或特殊待办。",
        "渠道支持站内信、邮件、短信、企业微信或微信服务通知，并记录送达与阅读状态。",
    ])
    add_heading(doc, "门户与内部系统集成", 2)
    add_table(
        doc,
        ["门户功能", "内部来源/回写"],
        [
            ("供应商询价报价", "ERP/SRM询价、价格库、比价、定标和采购订单"),
            ("采购/委外订单", "ERP订单发布，供应商确认、承诺交期和进度回写"),
            ("ASN与送货单", "SRM生成，WMS扫描收货，QMS检验，ERP入库"),
            ("供应商质量与8D", "QMS异常发布，供应商回复、整改和验证回写"),
            ("供应商对账与发票", "ERP采购/委外对账、应付和付款状态"),
            ("客户项目进度", "ERP订单、APS计划、MES进度、QMS状态经业务审核后发布"),
            ("客户交货与文件", "ERP/WMS发货、财务/外贸文件、QMS质量文件"),
            ("客户客诉与退货", "门户提交后创建CRM/RMA/QMS异常和返工补货流程"),
            ("香港客户订单", "沿用ERP普通客户订单、交付和报关资料状态；价格转换只在内部订单审核与授权报表中展示"),
        ],
        [2500, 6860],
        font_size=9,
    )


def section_bi(doc):
    add_heading(doc, "数据分析、指标与经营驾驶舱", 1)
    add_heading(doc, "管理层驾驶舱", 2)
    add_bullets(doc, [
        "销售收入、报价成交率、订单毛利、准交率、延期金额、库存周转、现金流和应收账龄。",
        "材料齐套、CNC/车床/后工序/组装产能与在制、计划达成、一次合格、返工、报废、OEE和质量损失。",
        "采购/委外准交、来料合格、供应商风险、呆滞金额、成本差异和部门经营。",
        "在岗人数、出勤、加班、缺卡、考勤与报工差异、劳务人数/工时/费用和结算进度。",
    ])
    add_heading(doc, "指标口径", 2)
    kpi_rows = [
        ("订单准交率", "按承诺交期完成交付的订单行数/应交订单行数", "订单、发货"),
        ("齐套率", "数量齐套、物料种类齐套、关键料齐套分别计算", "BOM、库存、在制、在途"),
        ("一次合格率", "首次检验合格数量/首次检验数量，不含返工后合格", "MES、QMS"),
        ("最终良率", "最终合格数量/投入数量", "MES、QMS"),
        ("返工率", "进入返工的数量/投入数量", "QMS、MES"),
        ("报废率", "报废数量/投入数量；同时按金额统计", "MES、成本"),
        ("材料利用率", "成品净用量/实际材料投入量", "领料、余料、报废"),
        ("计划达成率", "按期完成计划数量/计划数量", "APS、MES"),
        ("OEE", "时间开动率×性能开动率×合格率", "设备、MES、QMS"),
        ("采购准交率", "供应商按承诺日期足量交付的订单行/应交订单行", "采购、ASN、收货"),
        ("报价偏差", "实际成本-报价成本，并按材料/工时/委外/质量拆分", "报价、成本"),
        ("呆滞库存率", "按约定库龄及无未来需求库存金额/总库存金额", "库存、需求"),
        ("首检/巡检/末件按时完成率", "在规定时间内完成的对应检验任务/应完成任务", "QMS、MES"),
        ("考勤确认及时率", "在规定期限内完成确认的人员/应确认人员", "考勤、审批"),
        ("劳务结算及时率", "按结算日完成行政/部门/财务确认的劳务单位/应结算劳务单位", "考勤、劳务、财务"),
    ]
    add_table(doc, ["指标", "定义", "主要数据源"], kpi_rows, [1900, 5000, 2460], font_size=8.8)
    add_heading(doc, "分析钻取和根因", 2)
    add_bullets(doc, [
        "任何汇总指标都应下钻到客户、订单、产品、工单、工序、设备、人员、材料批次、供应商和凭证。",
        "异常原因统一编码，支持Pareto、趋势、同比、环比、目标差异和责任部门分析。",
        "报表口径、公式、负责人、刷新频率和变更历史纳入指标字典。",
        "财务和老板AI问答必须复用本指标字典与钻取链路，不得在问答中另行生成无法与正式报表勾稽的指标口径。",
    ])


def section_integration(doc):
    add_heading(doc, "集成、数据、安全与非功能要求", 1)
    add_heading(doc, "接口与事件集成", 2)
    add_bullets(doc, [
        "PLM向ERP发布物料、BOM、工艺和变更；ERP向PLM反馈库存、采购价和成本参考。",
        "ERP向APS提供订单、BOM、工艺、库存和资源；APS向MES下发计划；MES向ERP回传报工、耗料和完工。",
        "SRM/客户门户与ERP、WMS、QMS和财务实时或准实时交换订单、进度、文件、质量和结算。",
        "设备平台通过工业协议采集状态、产量、节拍、报警和能耗；关键数据带设备时间戳和来源。",
        "行政系统从考勤机、门禁或移动端采集原始打卡，与人员、班次、请假加班和MES报工关联；审核后的劳务结算向财务生成应付依据。",
        "接口具备唯一业务键、重试、幂等、错误队列、监控和对账，避免重复订单或漏传。",
    ])
    add_heading(doc, "权限、审计与职责分离", 2)
    add_bullets(doc, [
        "采用角色、组织、数据范围、字段和操作权限，敏感价格、成本、工资、供应商和客户信息按岗位隔离。",
        "供应商报价、定标、采购订单、收货、对账和付款职责分离；报价截止前可配置密封。",
        "所有主数据、单据、审批、文件、接口、登录、下载和库存调整保留审计日志。",
        "关键删除采用停用或冲销，不直接物理删除历史交易。",
    ])
    add_heading(doc, "AI数据接口与技术控制", 2)
    add_bullets(doc, [
        "AI查询通过白名单指标服务、授权数据视图或受控API完成，不允许模型直接持有生产数据库账号或执行任意SQL写操作。",
        "用户身份、组织、数据范围和字段权限在查询执行层再次校验，不能只依赖模型提示词约束。",
        "结构化查询、指标公式、参数、返回行数、引用单据和回答文本使用同一追踪编号，便于重放和核对。",
        "对上传文件和知识库内容进行来源、版本、恶意指令、敏感信息和访问权限检查，外部文档不能改变系统规则。",
        "建立模型可用性、响应时长、失败率、答案反馈、数字一致率和越权拦截监控；模型不可用时仍可使用标准驾驶舱和报表。",
        "模型、嵌入、提示模板和指标语义层升级前进行回归测试，固定财务问题集不得出现口径漂移。",
    ])
    add_heading(doc, "性能、可用性和移动端", 2)
    add_bullets(doc, [
        "车间扫码、仓库收发和门户查询应满足现场并发和响应要求；离线或网络中断需有明确降级与补传机制。",
        "支持网页、平板、手机和工业终端，关键操作界面按角色简化。",
        "明确可用性目标、备份频率、恢复时间、恢复点、容灾、监控、日志保留和升级窗口。",
    ])
    add_heading(doc, "数据迁移与切换", 2)
    add_bullets(doc, [
        "迁移组织部门、岗位、人员、班次、劳务单位、客户、供应商、物料、BOM、工艺、库存、在制、未结订单、应收应付、固定资产和必要历史价格。",
        "至少进行两轮模拟迁移和对账；切换前冻结规则、盘点、期初核对和未结业务清理。",
        "历史Excel应区分必须结构化迁移、只归档查询和不迁移三类。",
    ])


def section_vendor(doc):
    add_heading(doc, "产品路线与厂商选型建议", 1)
    add_callout(doc, "选型边界", "以下结论用于形成候选路线，不等同于厂商承诺。所有功能必须以具体产品版本、模块清单、接口方案、实施范围和真实场景演示为准。")
    add_table(
        doc,
        ["方案", "适配重点", "建议组合", "重点验证"],
        [
            ("鼎捷", "离散制造、工艺、APS、MES、QMS和WMS一体化路线", "E10/T100＋PLM＋APS＋sMES＋sQMS＋WMS＋SRM/门户＋BI/AI", "两类跨序返工、详细财务、AI口径引用、报价模型、套料合批和门户深度"),
            ("用友", "多组织、项目制造、业财和精细成本", "U9 cloud＋PLM＋MES/QMS/WMS＋供应商/客户门户＋BI/AI", "两类跨序返工、现场颗粒度、APS约束、AI权限和定制升级"),
            ("金蝶", "云化、财务供应链和制造均衡、流程配置", "云·星空＋PLM＋制造云/智慧工厂＋WMS/QMS＋SRM/门户＋BI/AI", "两类跨序返工、设备采集、机加工报价、余料和AI数据勾稽"),
            ("黑湖小工单", "中小工厂快速上线工单、扫码、进度、物料和轻质量", "小工单＋完整财务ERP＋供应链/门户＋BI/AI或定制集成", "完整财务成本、两类跨序返工、MRP/APS、外贸、对账和AI权限"),
        ],
        [1200, 2500, 3000, 2660],
        font_size=8.5,
    )
    add_heading(doc, "标准、增购与专项定制边界", 2)
    add_table(
        doc,
        ["类型", "典型内容"],
        [
            ("ERP标准/配置", "销售、采购、库存、财务、基础生产、MRP、应收应付、标准报表、审批和自动凭证"),
            ("同厂商增购模块", "PLM、APS、MES、QMS、WMS、SRM、设备、行政人事/考勤、BI、预算、合并和门户"),
            ("专项配置/定制", "两类跨序返工、CNC/车床、后工序、独立组装、FAI、劳务结算、机加工报价、订单合批、套料、余料、AI问答和特殊报表"),
            ("第三方接口", "CAD、单一窗口/报关行、物流、电子签章、银行、税务、交易所/持牌金属行情、考勤机/门禁、设备、AGV/立库、短信/微信和AI模型"),
        ],
        [2100, 7260],
        font_size=9.2,
    )
    add_heading(doc, "建议评分结构", 2)
    add_table(
        doc,
        ["评价维度", "建议权重", "评价方法"],
        [
            ("业务功能适配", "25%", "真实订单端到端演示和需求矩阵逐项响应"),
            ("制造与现场深度", "20%", "MRP/APS、扫码、返工、质量、仓储和设备场景"),
            ("财务成本与报表", "15%", "自动凭证、对账、工单成本、报价偏差和系统生成文件"),
            ("AI与数据可信度", "10%", "固定问题准确率、指标口径、单据引用、权限隔离、审计和模型治理"),
            ("门户与协同", "10%", "供应商报价/送货/对账及客户项目/文件"),
            ("技术与集成", "10%", "接口、扩展、性能、安全、移动端和升级"),
            ("实施与行业经验", "5%", "顾问、方法、数据治理、案例、计划和风险"),
            ("全生命周期成本", "5%", "许可/订阅、实施、定制、接口、运维、升级和扩容"),
        ],
        [2100, 1500, 5760],
        font_size=9.2,
    )


def section_implementation(doc):
    add_heading(doc, "项目治理、实施、培训与验收", 1)
    add_heading(doc, "项目组织", 2)
    add_table(
        doc,
        ["角色", "责任"],
        [
            ("项目委员会", "范围、预算、重大决策、跨部门资源和阶段验收"),
            ("业务项目经理", "进度、问题、风险、决策、供应商协调和成果提交"),
            ("IT项目经理", "架构、环境、接口、数据、安全、测试和运维"),
            ("模块负责人", "销售、工程、PMC、采购、生产、品质、仓库、财务、外贸的流程与验收"),
            ("关键用户", "数据准备、配置确认、测试、培训和一线推广"),
            ("实施供应商", "蓝图、配置、开发、迁移、测试、培训、上线和缺陷整改"),
        ],
        [2200, 7160],
        font_size=9.2,
    )
    add_heading(doc, "测试与上线", 2)
    add_bullets(doc, [
        "单元测试验证字段、规则、权限和打印；集成测试验证系统间数据；用户验收采用真实订单。",
        "对库存、在制、采购、委外、应收应付、成本和财务期初进行双向对账。",
        "上线采用明确的冻结时间、盘点、数据迁移、未结单处理、回退计划和现场支持。",
        "上线后建立问题分级、响应SLA、每日复盘、关键指标监控和阶段稳定标准。",
    ])
    add_heading(doc, "培训与变更管理", 2)
    add_bullets(doc, [
        "按角色提供流程培训、岗位操作、异常处理、报表使用和权限责任，结合实际工单练习。",
        "关键用户完成培训和考试后再培训一线人员；现场张贴扫码、返工、质量和仓库SOP。",
        "禁止系统上线后继续维护平行Excel；确有过渡需求须明确结束日期和对账责任。",
        "按月复盘准交、库存、报工、质量、成本和使用率，持续清理绕行流程。",
    ])


def section_acceptance(doc):
    add_heading(doc, "端到端演示与验收场景", 1)
    scenarios = [
        ("模具、样品与正式业务订单收费控制", ["分别建立一张收费模具订单、一张免费样品订单和一张正式业务订单。", "免费样品填写费用承担部门、预计成本、免费原因、关联商机和预计转正式订单日期，并完成业务经理、财务和授权管理层审批。", "正式业务订单将单价设为零或删除付款条件，验证系统阻断；恢复有效价格、核价、原始客户订单和付款条件后才允许提交。", "执行领料、报工、委外、检验和包装，证明收费或免费不影响实际成本归集。", "把样品转为正式订单并执行样品费抵扣，证明原样品订单、成本、审批和抵扣履历不被覆盖；输出三类订单收费与成本报告。"]),
        ("新客户报价到订单时效监测T0", ["新客户询价时查看实时金属材料价格、历史相似产品报价/成交/实际成本并形成核价快照。", "客户确认后由业务建立客户主数据并提交BOM申请，工程完成BOM、工艺和内部图纸转换。", "模拟新程序尚未完成：系统允许BOM完成通知和业务下单，但阻止对应CNC/车床工序发放或开工，直至程序验证发布。", "业务订单提交业务经理时生成T0，验证主流程总时钟和各节点计划同时启动。", "分别完成业务经理审核、财务审核和跨部门订单评审，证明缺核价或缺原始客户订单时财务必须阻断。"]),
        ("订单评审到MRP、采购和到料全链时效", ["订单评审结束后模拟PMC未及时跑MRP，验证临期、超期和升级预警。", "运行MRP并保存运行批次、输入快照、净需求解释、异常处理及确认发布时间。", "从采购建议发布开始记录采购领取、询价/定标、采购单提交和财务审核时长。", "记录供应商确认、备货、发运、运输、到厂、仓库收料和IQC放行的独立起止时间。", "系统输出每个节点的首次响应、排队、处理、暂停、总历时、P90/P95、按时率和超期责任。"]),
        ("逐工序品质门禁与数量差异交接", ["CNC完工申报100件，品质只对送检对象做质量判定，不确认数量。", "CNC发起转移100件，打磨实点94件，系统先转移94件并将差异6件继续留在CNC。", "系统向CNC责任人和主管发出数量差异预警，并证明94件可以按受控路线继续生产。", "差异6件使用数量差异单关闭；另模拟6件质量不合格，证明必须使用NCR/MRB而不能与物理短少互相冲销。", "逐道验证品质确认、数量交接、打磨、委外表处、回料、镭雕、丝印、全检和包装的时间戳及门禁。"]),
        ("客户报价到订单", ["上传客户图纸并保存版本。", "检索相似产品、采购和委外历史价。", "填写材料、工时、工装、委外、包装和毛利，生成成本分析表。", "完成工程/财务/管理审批并生成客户报价PDF。", "报价转订单并触发合同评审。"]),
        ("实时金属材料价格到报价与成本", ["接入一条许可行情并映射到企业材料牌号、单位和币种。", "叠加汇率、升贴水、加工、物流和税费计算企业落地参考价。", "采购查看供应商与市场偏离，财务模拟库存/在制/毛利影响，业务仅查看授权参考价。", "生成客户报价并保存行情、汇率和公式快照。", "模拟行情超过阈值，系统触发重新核价预警但不自动改写已审批采购订单或法定存货账。"]),
        ("PMC滚动齐套", ["按订单和BOM展开需求。", "扣除已分配、冻结和不良库存，加入按期到货和在制。", "计算净需求和建议采购/生产量。", "模拟采购到货、报废和订单优先级变化，齐套表实时更新。"]),
        ("每一道工艺独立成本核算", PROCESS_COST_ACCEPTANCE),
        ("十三部门工作台与数据权限", ["分别以财务、总经办、业务、工程、PMC、采购、委外、仓库、生产、品质、后工序、行政和IT账号登录。", "验证各部门首页只显示本部门任务、待办、指标和授权数据。", "总经办可分析全厂但不能修改凭证、检验、库存和考勤原始记录。", "采购与委外、生产与后工序、后工序与组装的任务和统计可以区分。", "跨部门交接通过正式单据、状态、责任人和审计完成。"]),
        ("CNC/车床到后工序与独立组装", ["CNC或车床使用正确程序、材料、设备和人员扫码开工并通过首检。", "工件按路线进入清洗、丝印/镭雕、打磨/振磨并记录各自参数。", "需要组装的产品生成独立组装BOM、齐套、领料和组装工单。", "记录组件批次/序列号、扭矩/参数、测试和组件谱系。", "完成出货前100%全检、包装、标签和OQC后才允许发货。"]),
        ("首检、巡检、末件与FAI电子质量档案", ["新样品工单根据图纸和气泡号生成FAI任务。", "换型开工触发首检，生产到指定数量/时间触发巡检。", "工单结束前触发末件检，模拟末件超差并反查影响数量。", "检验记录保存原始测量值、检具、设备、人员、程序、图片和电子签名。", "审核后生成FAI及首检/巡检/末件报告，并证明旧版本不可被当前数据覆盖。"]),
        ("行政考勤与多劳务结算", ["导入两个劳务单位和正式员工的原始打卡、班次、请假和加班。", "系统计算个人、班组、部门和劳务单位的日报/月报及异常。", "将考勤与MES报工比较，生成有打卡无报工和有报工无打卡差异。", "按两个劳务单位不同费率、加班、补贴、扣款和管理费生成对账与结算单。", "行政、用工部门和财务分级确认后转应付，并验证普通用户不能查看个人敏感明细。"]),
        ("内部跨序返工路线A：返工后直接返回第五工序", ["第五工序发现第二工序问题并上传图片。", "冻结受影响数量并确认第三、第四工序原结果仍然有效。", "生成第五→第二返工→复检→第五的受控路线。", "第二工序返工复检合格后直接返回第五工序。", "第三、第四工序原履历保留且有有效性确认，返工成本和交期同步变化。"]),
        ("内部跨序返工路线B：返工后从第三工序顺序重流", ["第五工序发现第二工序问题并上传图片。", "质量与工程判定第二工序返工会影响第三、第四工序结果。", "生成第五→第二返工→第三→第四→第五的受控路线。", "系统阻止从第二工序直接跳回第五工序，每道重流工序均需报工和复检。", "第三、第四工序原履历标记被替代但不删除，重新归集工时、成本、良率和交期。"]),
        ("跨序返工同批数量拆分", ["从同一异常批次选择不同受影响数量。", "部分数量指定路线A，部分数量指定路线B。", "系统自动拆分子批次和二维码，分别显示在制位置与路线。", "未满足各自复检和汇合条件前禁止合批。", "关闭后汇总原批次、子批次、成本和质量履历。"]),
        ("包装发现异常、CNC责任的返工成本与二次委外应付", ["在包装发现不合格，冻结受影响数量，由品质创建NCR并初判责任为CNC。", "系统从原生产订单拆出独立返工子订单，建立包装→CNC返工→品质复检→打磨去毛刺→品质→委外表处→回料品质→镭雕→品质→丝印→品质→全检→重新包装的受控路线。", "CNC、打磨、品质、镭雕、丝印、全检和包装的第二次人工、机时、耗材与制造费用均以新履历重新计入返工成本，不能复用原订单第一次成本。", "系统生成新的付费委外表处订单；因表处供应商并非责任方，其第二次合格加工应正常进入委外对账与应付，且与原委外应付分开。", "返工毛成本整体归集到CNC责任维度，原订单成本、返工新增毛成本、回收金额和企业净损失分别显示；责任复核变化时以新版本重分类但不删除执行事实。"]),
        ("包装发现异常、委外表处供应商责任的无偿返工与质量损失", ["在包装发现不合格，冻结受影响数量，由品质创建NCR并认定责任为委外表处供应商。", "系统拆出独立返工子订单，并生成供应商责任无偿返工/退回单，而不是新的付费委外订单。", "供应商本次返工的追加应付为零，系统阻止将该无偿返工重复带入委外对账；采购/委外仍可对退运、回料、补货、8D、索赔和时效进行跟踪。", "企业发生的退运物流、内部搬运、复检、重新镭雕/丝印/全检/包装、报废、停线、交期延误或客户索赔等实际成本仍逐项计入返工毛成本，并全部按责任比例归集为该供应商造成的质量损失。", "报告同时显示市场等价返工费（管理分析、不进入法定实际应付）、实际追加应付、供应商造成的毛损失、可追偿、已扣款/已收回和企业净损失。"]),
        ("香港代生产客户正式订单70%价格", ["在客户主数据勾选一个特定客户为香港代生产价格客户，分别建立正式、样品和模具订单并输入相同原始单价。", "证明只有该客户的正式业务订单自动生成原始价格×70%的订单价格，样品和模具订单均保持原始价格；再建立未勾选客户正式订单，证明同样保持原价。", "切换客户、切换订单类型并修改原始单价，验证系统自动重判、重算和生成新版本；尝试手工覆盖计算价或重复乘70%时阻断并预警。", "让业务经理和财务按普通订单流程审核，随后完成订单评审、MRP、采购、生产、品质、委外、包装、报关、开票和收款，证明除价格换算外没有增加或删减任何流程。", "用计算后正式订单价格与全部材料、人工、机时、委外、检验、返工、包装及工序成本计算毛利，并输出价格计算审计表。"]),
        ("供应商在线报价到送货", ["发布带图纸和阶梯数量的RFQ。", "三家供应商网页报价并提交技术偏离。", "多轮议价、保留版本、生成比价和定标。", "中标转采购/委外订单，供应商确认交期。", "供应商创建ASN并打印送货单和标签，仓库扫码收货。"]),
        ("供应商质量与对账", ["IQC或委外检验不合格。", "供应商门户收到异常并提交8D/补货。", "检验和整改关闭后更新绩效。", "系统生成采购/委外对账，供应商在线确认并上传发票。"]),
        ("客户门户", ["客户仅能看到自己的订单和受控里程碑。", "业务发布交货计划，客户在线确认。", "发货后客户下载装箱单、发票和检验报告。", "客户提交客诉并查看返工/补货进度。"]),
        ("工序与审批预警", ["模拟工序未按时开工、进度落后和已超期。", "模拟报价、合同或返工审批超时。", "系统按级别提醒、升级、记录原因和预计恢复。", "完成后验证关闭并进入超期分析。"]),
        ("财务AI问答", ["询问未来30天资金缺口并显示回款、付款和假设。", "询问某订单实际成本高于报价的原因并拆解材料、人工、机时、委外、返工和报废。", "询问金属行情变化对在途采购、库存、在制、标准成本和订单毛利的模拟影响，并引用行情快照。", "询问月结未完成任务和业务财务对账差异并下钻原单。", "核对答案与正式报表数字、口径和数据截止时间一致。", "验证AI不能直接记账、付款、关账或越权查看工资和银行明细。"]),
        ("老板AI经营问答", ["询问本月订单、销售、回款、毛利、现金、库存、交付和质量与目标及同期差异。", "继续追问毛利下降原因并下钻客户、产品、订单、工序、供应商和质量损失。", "模拟材料涨价或设备停机对交期、毛利和现金的影响，并明确标记为情景模拟。", "答案提供指标定义、查询范围、引用报表和原始单据。", "使用无权限账号验证敏感成本、供应商底价和其他组织数据不可见。"]),
        ("系统生成文件", ["基于同一订单自动生成报价、成本、订单/工序/采购/齐套/供应商进度、质量、FAI、后工序、组装、全检包装、考勤、劳务结算、送货单、对账和分析报表。", "修改源数据并重新审批后生成V2，V1保留且标记失效。", "报表可钻取原始单据，导出Excel/PDF不需要人工拼接。"]),
    ]
    for title, steps in scenarios:
        add_heading(doc, title, 2)
        add_numbered(doc, steps)
    add_callout(doc, "否决条件", "如果关键报表仍需复制到Excel再加工，组装只能作为包装备注、后工序只能一次性报完，每道实际工艺不能建立独立成本卡、工序成本不能从本工序增量累计到订单总成本、返工重复执行不能再次计费，首检/巡检/末件/FAI不能保存原始测量与版本，考勤劳务仍靠人工汇总，跨序返工不能支持两类路线，返工不能按子订单重算全部重复工序成本、不能区分正常二次委外应付与供应商责任无偿返工、不能核算供应商造成的毛损失与净损失，不能做到仅已勾选香港代生产价格客户的正式业务订单按输入价格70%计价、样品模具及其他客户保持原价、价格以外流程完全不变，AI没有口径引用或能越权，或预警没有责任升级关闭，则不视为满足需求。", fill="FCECEC", color=RED)


REQ_GROUPS = {
    "GOV": ("项目治理", [
        "建立项目委员会、业务与IT项目经理及模块负责人", "建立范围、变更、问题、风险和决策台账", "按阶段设置业务和数据验收门槛", "真实订单作为UAT和选型演示数据", "禁止未经批准扩大首期范围", "明确上线冻结、回退和稳定期机制"
    ]),
    "ORG": ("组织与部门", [
        "公司工厂部门班组工作中心成本中心利润中心分层", "财务部总经办业务部工程部PMC采购委外仓库生产部品质部后工序部行政IT十三部门", "生产部CNC与车床工作中心分离", "后工序清洗丝印镭雕打磨振磨全检包装工作中心", "组装独立工作中心和独立组装工单", "采购与委外角色单据进度和对账分离", "员工主岗位兼任岗位和权限有效期", "人员技能设备和工序授权", "部门工作台待办任务超期异常和指标", "菜单操作组织行字段状态金额权限", "总经办全厂分析但不可越权修改源数据", "部门数据责任人和指标口径负责人", "跨部门正式单据状态责任人和审计交接", "请假审批代理和到期自动收权", "组织部门岗位工作中心变更保留历史", "部门工作中心成本和产能映射", "设备职能即使不独立成部门也有角色", "敏感财务成本价格考勤和个人信息字段隔离", "部门绩效可钻取源单但不改变法定账", "十三部门账号权限和跨部门交接UAT"
    ]),
    "OTL": ("订单全生命周期与时效", [
        "报价查询实时金属材料价格和历史相似产品报价", "客户确认后业务新建客户并提交BOM申请", "工程建立BOM工艺和内部图纸", "程序任务可与BOM并行但在机加工开工前必须发布", "BOM完成状态和程序可开工状态分别管理", "业务订单提交时建立T0主流程时钟", "业务订单必须关联核价和原始客户订单", "业务经理审核商业条款和交期承诺", "财务审核核价原始订单信用毛利和付款条件", "跨部门订单评审按专业会签和条件项关闭", "评审通过后PMC运行确认并发布MRP", "MRP建议到采购领取下单全过程计时", "采购订单财务审核和变更重审", "供应商确认备货发运运输到厂分段计时", "仓库收料和IQC检料分责计时", "切料按工艺条件启用或记录不适用", "CNC车床每道工序独立开暂停完工计时", "每个受控工序之间设置品质确认门禁", "品质只确认质量不确认交接数量", "交出方申报接收方实点并取较小数实转", "CNC报100打磨收94时先转94并向CNC预警6", "数量差异单与质量NCR严格分离", "打磨品质外发回料镭雕丝印全检包装逐节点计时", "委外等待加工检验运输回料分段计时", "订单T1为全检合格包装完成形成可发货批次", "交付后尾数返工入库直接入库或报废四路径", "订单结案前投入交付入库报废在制数量强制平衡", "基线计划最新承诺实际时间并存且改期留痕", "首次响应排队处理暂停总历时超期分别统计", "退回重提保留每轮和总历时", "节点超期重算后续预计完成与客户交期风险", "输出平均中位P90P95按时率超期率和积压", "生成订单全生命周期时间轴和瓶颈Pareto", "订单类型与收费方式分别控制", "模具样品允许收费免费或部分收费但必须审批", "正式业务订单必须收费且零价阻断", "免费订单仍完整归集实际制造成本", "模具样品转正式订单保留来源和抵扣分摊履历"
    ]),
    "MDM": ("主数据", [
        "物料编码唯一并支持重复校验", "原材料长宽高/直径和成品长宽高/直径分字段", "毛坯重量、成品重量、密度和单位换算", "多层BOM替代料客户定制和版本", "独立组装BOM组件谱系和版本", "完整工艺路线覆盖CNC车床后工序组装全检包装", "CNC车床程序按产品工序机型版本管理", "程序发布下载修改使用和校验日志", "图纸在线预览版本水印和下载日志", "工程变更影响程序检验库存采购在制委外", "样品订单试制工单和打样记录", "工装刀具模具检具台账与寿命", "设备产品工序能力矩阵", "人员技能和设备授权", "组织部门岗位班组工作中心主数据", "班次日历用工类型和劳务单位主数据", "仓库库区库位和摆放标准", "客户供应商主数据及信用质量属性", "财务科目成本中心和凭证规则", "主数据新增修改停用审批", "主数据质量报表"
    ]),
    "QTN": ("报价与合同", [
        "客户询价和图纸统一归档", "样品批量备料报价分类", "历史报价和成交价查询", "相似产品特征搜索", "历史采购和委外价格引用", "实时金属行情和企业参考价引用", "报价保存材料行情汇率和价格快照", "行情超阈值或快照过期重新核价", "材料尺寸重量利用率和损耗成本", "工序机时人工刀具工装成本", "委外最低批量运输和损耗成本", "阶梯数量价格和一次性费用摊销", "报价毛利和风险系数", "报价多级审批", "中英文报价PDF系统生成", "报价版本和有效期", "报价转订单不重复录入", "合同扫描和纸质档位置", "跨部门合同评审", "订单变更影响分析", "客户退货返工和补货", "模具样品正式业务订单三分类", "模具样品收费免费部分收费和抵扣分摊", "正式业务订单强制收费和付款条件"
    ]),
    "HKO": ("香港代生产客户正式订单价格", [
        "客户主数据香港代生产价格客户勾选项", "勾选启用停用申请审核生效和变更日志", "仅勾选客户且正式业务订单同时成立时触发", "业务输入原始单价金额完整保存", "正式订单价格自动等于原始输入价格百分之七十", "样品订单和模具订单保持原输入价格", "未勾选客户所有订单保持原输入价格", "客户勾选状态和订单类型建立下单时快照", "原始价格系数计算后价格和舍入版本可追溯", "禁止业务手工修改固定系数和计算后价格", "防止业务先手工乘七十后系统重复折算", "客户订单类型原价数量币种变化自动重判重算", "已审核订单价格变化走普通变更和重新审核", "业务经理和财务沿用普通订单审核权限", "BOM评审MRP采购生产品质委外包装流程不变", "报关发货开票应收收款总账流程不变", "不建立独立三十分留存或特殊分成结算流程", "计算后正式订单价格用于应收成本毛利经营分析", "每道工艺和全部实际制造成本继续完整归集", "价格漏算误算重复折算和版本不一致预警", "香港客户正式订单价格明细与审计报告", "样品模具及其他客户未应用七十的反向验证"
    ]),
    "PMC": ("PMC/MRP/APS", [
        "MRP考虑订单BOM库存在制在途损耗和安全库存", "合格待检冻结不良和已分配库存分离", "采购生产委外调拨建议", "MRP建议来源可追溯", "建议单经审批转正式订单", "动态成品备料和齐套表", "数量品种和关键料齐套率", "订单库存报工报废冻结到货变化实时重算", "周交付周生产和日机台计划", "客户交期和订单优先级", "机台人员夹具模具检具共用资源冲突", "有限产能和设备负荷", "插单和交期模拟", "替代设备外发加班建议", "相同材料工艺机台订单合批提示", "余料匹配和套料优化", "合批节料与交期影响比较", "计划责任人和预计恢复", "采购委外供应商工序进度联动", "平衡化生产和在制控制", "自动生成初版架机表并经PMC确认下发", "架机表调整留痕冲突检查和架机率统计", "返工子订单纳入独立产能负荷与交付重排", "重复委外和后续工序重新计算齐套与预计完成"
    ]),
    "SRM": ("采购与供应商", [
        "供应商网页结构化报价", "原料标准件定制件委外工装设备多类询价", "数量阶梯价格和多报价单位", "报价产能交期付款和技术偏离", "供应商报价与市场基准和历史成交比较", "多轮议价和报价版本", "密封报价和截止控制", "供应商报价PDF和自动比价", "综合成本质量交期绩效评审", "多部门定标和多源分配", "中标报价转采购/委外订单", "供应商在线确认订单和拆分交货", "供应商维护采购/委外进度", "ASN预到货", "网页打印送货单和标签", "送货数量版本质量文件校验", "供应商质量异常和8D", "采购对账在线确认", "委外对账在线确认", "发票上传和付款查询", "供应商数据隔离和图纸水印", "区分正常二次付费委外订单与供应商责任无偿返工单", "供应商责任无偿返工追加应付为零且禁止重复对账", "供应商质量损失索赔扣款回收与净损失跟踪"
    ]),
    "MKT": ("金属材料价格中心", [
        "铜铝钢材不锈钢镍锌锡铅钛和自定义合金", "接入交易所官方或持牌行情数据", "供应商最新报价和数量阶梯", "内部历史采购成交价格曲线", "行情品种合约牌号形态规格映射", "现货期货交割月买卖结算价", "币种单位时区行情时间和接收时间", "实时延时日结人工价状态标识", "市场基准汇率升贴水加工物流税费落地价", "采购查看供应商明细和市场偏离", "财务查看成本库存在制承诺和毛利模拟", "业务查看授权参考价和报价影响", "业务隐藏供应商身份底价和议价记录", "采购订单保存行情汇率和报价快照", "已审批订单不被行情自动改写", "标准成本重估先模拟后审批", "报价快照过期和涨跌阈值重新核价", "材料价格涨跌和供应商偏离预警", "行情延迟中断异常值和许可到期预警", "接口认证限流重试备用源和交易日历", "原始行情标准化公式修正和用途审计", "历史时点报价成本可重现", "交易所数据内部使用展示定价和再分发许可", "未经许可禁止网页抓取和对外传播"
    ]),
    "MES": ("生产与MES", [
        "工单按交付批次和工艺拆分", "CNC车床工作中心设备组班次和能力分离", "派工校验材料图纸程序工装检具", "原料批次工单批次单件一码到底", "扫码人员设备工单开工暂停完工", "投入良品不良返工报废数量", "电子图纸指导书和程序版本", "上道未完成不得下道开工", "首件未合格不得批量生产", "条件跳工序审批和日志", "中途改工序经工程变更生成新路线版本并限定适用批次", "清洗参数批次温度时间和检查", "丝印图稿网版油墨固化和首件", "镭雕程序内容序列号参数和校验", "打磨磨料工具时间和外观尺寸", "振磨介质研磨液装载时间和防混", "出货前百分百全检包装标签和照片", "组装独立BOM齐套领料和工单", "组装组件批次序列号谱系和扭矩参数", "组装过程检查功能测试返工和成本", "同工序返工任务", "跨工序退回和专用返工路线", "跨序路线A返工后直接返回发现工序", "跨序路线B返工后从下一工序顺序重流", "同批数量拆分子批次和独立二维码", "重流工序原履历标记被替代但不删除", "返工路线顺序约束和汇合条件", "返工复检和质量放行", "返工路线重算在制负荷交期和成本", "原工序履历不可删除", "设备人员产品能力校验", "机台日任务和人员日任务", "机台和人员良率达成率", "计时计件工资来源于审核报工", "设备点检保养维修和备件", "设备状态与APS联动", "停机和异常闭环", "OEE和设备绩效", "每次返工建立独立子订单并关联原订单NCR和责任版本", "返工重新经过的每道内部委外检验包装工序重新报工计费", "返工执行事实与责任成本归属分离保存"
    ]),
    "OPC": ("工序级成本", [
        "成本核算到订单工单批次路线版本工序执行实例", "每道实际工艺建立独立成本卡", "同一工序重复返工执行建立新实例", "工序保存标准预计实际暂估实际结算四口径", "工序成本口径费率版本币种截止时间和状态", "上道累计成本随合格数量转入下道", "本工序新增成本与转入累计成本分栏", "直接材料按工序领退补料和批次归集", "人工按准备调机加工上下料检验包装时长归集", "机器按设备运行准备时长和有效费率归集", "刀具夹具模具治具寿命与摊销归集", "委外每种工艺加工费运费损耗返工扣款独立归集", "首检巡检末件FAI复检全检OQC成本独立归集", "工序间搬运容器标签包装和物流成本归集", "制造费用按批准动因和费率版本分配", "正常损耗超额损耗返工报废停机等待分栏", "供应商索赔残值责任扣款和回收分栏", "投入合格不良返工报废在制转出数量平衡", "品质判定数量交接和财务计价职责分离", "本工序新增成本和新增单位成本计算", "工序完工累计成本和累计单位成本计算", "工序报废承接累计成本和残值计算", "工序标准实际差异按材料人工机器委外良率拆解", "切料CNC车床打磨清洗振磨分别成本", "委外表处镭雕丝印分别成本", "组装检验全检包装分别成本", "部分完工等价产量和跨月在制分配", "拆批合批按批准动因分配且留底稿", "迟到发票和费率变化受控重算生成新版本", "工序成本结算和月结锁定", "无费率零负成本数量不平成本链断裂预警", "未结委外漏报工费用未分配月结阻断", "工序成本明细卡和订单成本瀑布", "工序标准实际差异和单位成本良率分析", "在制工序成本返工报废和完整性报告"
    ]),
    "QMS": ("品质与QMS", [
        "IQC首件巡检末件完工和出货检验", "新产品打样重大变更和客户FAI", "FAI气泡图全尺寸材料性能外观和文件", "系统自动生成FAI受控报告", "工序检验点数标准上下限和检具", "自动生成正确版本检验任务", "首检换型换程序换刀材料返工触发", "巡检按时间数量班次风险自动触发", "末件检超差反查最近合格后的影响数量", "检验原始测量照片附件签名版本和日志", "出货前全检与包装质量状态联动", "质量异常上传照片和测量数据", "不合格批次自动冻结", "同批同设备同材料影响检查", "MRB返工返修退回让步降级报废", "MRB判定直接返回或顺序重流策略", "记录受影响特性及有效失效工序", "让步和放行审批", "内部和客户退货返工统一闭环", "返工前后测量和复检", "一次合格与最终良率区分", "NCR CAPA 8D和重复问题", "检具校准临期和禁用", "供应商质量在线协同", "客户客诉RMA和返工进度", "质量成本归集", "正向和反向追溯", "品质异常分析和Pareto", "NCR记录发现工序责任方责任比例和证据", "责任初判终判版本与返工路线成本联动", "供应商无偿返工与内部重复工序质量成本完整归集"
    ]),
    "WMS": ("仓储与WMS", [
        "ASN扫码收货和订单核对", "待检合格冻结不良库存状态", "上架推荐和电子仓位图", "标准摆放容器容量和防混", "批次炉号序列号追溯", "余料尺寸重量批次和原工单", "新订单优先匹配余料", "生产备料拣配配送补退料", "紧急缺料和配送任务", "扫码校验物料库位工单数量", "年度和循环盘点", "库存差异审批和原因", "非仓库人员无出入库调整权限", "负库存和无单移动控制", "呆滞物料自动识别和处置"
    ]),
    "FIN": ("财务与成本", [
        "多组织账套科目辅助核算和会计期间", "总账凭证结转汇兑和期末处理", "客户信用预收应收账龄催收和回款", "订单发货发票应收收款核销链路", "采购委外收货对账发票应付付款链路", "供应商预付质保金扣款和多币种付款", "现金银行银企对账和资金计划", "未来7/30/90天现金流预测", "费用申请借款报销标准和预算校验", "固定资产低值易耗和租赁", "存货计价暂估和库存总账勾稽", "税码销项进项发票查重和红冲", "外贸币种汇率汇兑出口收汇退税资料", "部门项目科目预算占用执行和调剂", "业务单据自动生成凭证", "自动凭证异常队列和源单冲销", "业务库存往来资产成本总账自动对账", "行政劳务对账结算转应付和付款", "劳务人数工时费率管理费补贴扣款核对", "标准成本和实际成本", "金属行情变动对库存在制采购承诺和毛利模拟", "材料行情快照审批后形成新标准成本版本", "产品订单工单工序项目批次成本", "CNC车床后工序和组装成本分别归集", "材料人工机时制造费用委外归集", "完工在制约当量和跨月成本分配", "材料价差量差和人工设备效率差异", "跨序路线A和路线B返工成本分别归集", "重流工序历史成本调整和质量损失", "报废累计成本和回收价值", "质量成本四分类", "报价标准实际成本对比", "客户产品订单实际毛利", "部门工作中心设备利润中心核算", "月结任务前置条件超期预警和关账", "法定报表管理报表勾稽和电子会计档案", "采购委外对账转应付", "外贸币种汇率和出口文件", "模具样品免费减免成本承担和批准额度", "正式业务订单收费应收和回款", "模具押金返还和正式订单成本分摊", "返工子订单独立成本对象并与原订单汇总", "重复工序材料人工机时制造费用检验包装物流委外再次归集", "正常二次委外应付与供应商责任无偿返工零应付分开", "返工毛成本责任归集回收金额企业净损失四层核算", "供应商市场等价返工费仅作管理分析不得虚增法定应付", "勾选香港代生产价格客户的正式订单原价与百分之七十计算价并列保存", "样品模具和其他客户不应用百分之七十", "香港客户价格转换后沿用普通应收开票收款和总账流程"
    ]),
    "ADM": ("行政考勤与劳务", [
        "员工工号部门岗位班组用工类型劳务单位", "入离职合同有效期和考勤地点", "正常白夜跨日轮班和休息班次", "法定节假日调休和企业日历", "考勤机门禁移动打卡和批准文件导入", "原始打卡设备地点来源和导入批次不可覆盖", "应出勤实出勤正常加班和请假工时", "迟到早退旷工缺卡出差培训和补卡", "考勤异常员工主管确认申诉和审批", "个人班组部门用工类型劳务单位日报周报月报", "加班请假缺卡和出勤趋势报告", "多个劳务单位和合同费率版本", "按人数天数正常加班小时班次岗位和计件结算", "劳务管理费补贴扣款税费和发票", "已审核考勤和MES有效报工计件作为结算依据", "生成劳务对账单结算单和差异明细", "行政用工部门财务分级确认", "劳务结算转财务应付和付款", "考勤与MES有打卡无报工差异", "考勤与MES有报工无打卡差异", "跨部门借调和超长工时异常", "原始打卡调整审批封存和审计", "总经办汇总部门主管本部门个人隐私隔离", "劳务合同费率结算和考勤超期预警"
    ]),
    "DOC": ("单据与报表", [
        "统一模板语言印章编号和版本", "正式文件由源数据自动生成", "修改源数据重新审批生成新版本", "历史文件快照不可被当前数据改写", "客户报价单", "内部成本分析表", "外贸发票装箱单报关资料", "订单进度表", "各工序进度表", "采购进度表", "PMC备料齐套表", "供应商进度表", "供应商报价和比价表", "采购对账单", "委外对账单", "仓位图", "首检巡检末件检验报告", "FAI报告", "后工序履历与参数报告", "出货前全检与包装报告", "组装工单与组件谱系报告", "品质异常分析", "成本核算分析", "工序成本明细卡与成本瀑布", "工序标准实际成本差异分析", "在制品工序成本报告", "工序成本完整性与月结检查", "停机报废良率损耗分析", "仓库呆滞物料分析", "金属材料价格与影响分析", "考勤日报月报", "考勤与MES工时差异报告", "劳务对账与结算单", "跨工序返工路线与成本分析", "返工子订单成本明细表", "供应商质量损失与索赔报告", "香港客户正式订单价格计算与审计表", "AI问答分析报告", "三类订单收费成本回收和转化报告", "订单尾数清单与尾数入库占用履历", "未完结工单账龄表", "架机表与架机率报表"
    ]),
    "ALT": ("预警中心", [
        "蓝黄橙红深红五级预警", "工作日班次小时和紧急度时间规则", "工序未开工无报工进度落后和超期", "预计影响客户交期", "报价核价和合同评审超期", "图纸BOM工艺程序工程变更审批超期", "采购委外确认到货和进度超期", "金属价格涨跌供应商偏离和快照过期", "行情延迟中断异常值和许可到期", "首件巡检末件FAI返工复检CAPA超期", "清洗丝印镭雕打磨振磨参数或进度异常", "组装缺料测试不合格和全检包装超期", "考勤缺卡未批加班和工时差异", "劳务合同费率对账和结算超期", "设备点检保养报修维修超期", "客户供应商待确认超期", "责任人主管部门负责人分级升级", "请假代理和转交", "原因措施预计恢复和关闭验证", "超期看板和处理效率", "预警不可由责任人删除", "长时间工单未完结账龄分层监控和疑似死单强制评审", "正式订单零价或缺付款条件阻断", "免费模具样品未审批或实际成本超额度", "返工子订单未建或重复工序漏报工漏计成本", "CNC责任二次表处未建正常应付委外单", "供应商责任无偿返工误生成追加应付或重复对账", "责任未终判索赔超期与企业净损失持续扩大", "工序成本模板费率缺失漏报工和零负成本", "工序数量成本不平成本链断裂和返工重复成本遗漏", "未结委外未分费用工序成本卡阻断月结", "香港客户正式订单百分之七十漏算或重复折算", "香港客户样品模具误算或其他客户误用百分之七十", "客户订单类型原价变化未重算和下游价格版本不一致"
    ]),
    "CUS": ("客户协同", [
        "客户网页登录", "客户数据和项目隔离", "查看报价合同订单和交货计划", "受控项目里程碑", "客户在线确认图纸样品和计划", "发货物流和签收", "下载报价装箱发票检验质量文件", "提交客诉退货补货", "查看返工补货关闭进度", "内部成本供应商人员信息隐藏", "文件水印版本和下载日志", "客户消息和超期待办", "香港代生产价格客户沿用普通客户门户和权限", "内部原始价格与百分之七十计算价按字段授权", "客户勾选不产生独立门户流程"
    ]),
    "BI": ("分析与指标", [
        "总经办全厂经营驾驶舱", "财务生产品质后工序工程PMC采购委外行政部门视图", "报价成交和未成交原因", "订单准交和延期原因", "齐套缺料和采购委外准交", "CNC车床后工序组装计划达成和在制", "首检巡检末件FAI按时完成率", "一次合格返工报废和最终良率", "材料利用率和超额损耗", "设备OEE停机原因", "供应商价格交期质量绩效", "库存周转和呆滞金额", "考勤出勤加班缺卡和工时差异", "劳务人数工时费用和结算进度", "报价标准实际成本偏差", "客户产品订单工单工序设备人员钻取", "指标字典和口径负责人", "趋势同比环比目标差异", "根因Pareto和重复问题", "模具样品正式订单结构收费成本回收和样品转化", "返工子订单毛成本回收净损失和责任部门供应商排名", "供应商追加应付与供应商造成质量损失分栏分析", "各工艺新增累计单位成本和成本占比", "工序标准实际差异在制成本和成本链完整性", "香港客户正式订单原价与百分之七十计算价分析", "香港客户计算后订单价格成本毛利回款和交付分析"
    ]),
    "AIX": ("AI分析问答", [
        "财务工作台和老板驾驶舱自然语言问答", "问答模型采用DeepSeek V4 Pro并记录模型版本用量", "连续追问筛选修改收藏和导出", "经审批指标语义层和只读查询服务", "受控会计政策制度指标字典知识库", "财务应收账龄回款和催收分析", "财务应付付款资金缺口分析", "订单报价标准实际成本差异分析", "客户产品订单毛利变化分析", "月结未结任务和对账异常检查", "发票税务预算费用异常分析", "金属行情供应商偏离和成本毛利影响分析", "老板经营总览和目标预算同期比较", "老板交付现金客户供应商风险追问", "材料汇率插单停机回款延迟情景模拟", "答案引用行情源时间许可单位和快照", "答案显示截止时间口径筛选计算和来源", "金额指标可钻取报表凭证和业务单据", "事实推断预测建议明确区分", "缺失数据明确说明不得编造", "继承组织行级字段级敏感权限", "默认只读不得自动记账付款关账审批", "写回仅生成草稿二次确认并走原审批", "问题模型数据引用导出和反馈全审计", "企业数据不用于外部模型训练和传输加密", "固定问题集与人工报表一致性验收", "权限越权提示和敏感数据防泄露测试", "工序成本占比标准实际差异和毛利影响问答", "工序成本可下钻材料人工设备委外检验和凭证", "香港客户正式订单原价百分之七十计算价成本毛利问答", "香港价格规则适用条件版本和异常引用"
    ]),
    "SEC": ("安全与非功能", [
        "角色组织数据字段操作权限", "十三部门及工作中心数据范围隔离", "客户供应商行级数据隔离", "敏感价格成本工资考勤个人信息隐藏", "总经办汇总和个人明细权限分离", "多因素认证和账号有效期", "兼任代理权限自动到期", "文件水印下载和登录日志", "接口幂等重试错误队列和监控", "网页平板手机工业终端", "网络中断降级和补传", "备份恢复容灾和日志保留", "主数据单据审批接口审计", "关键删除采用停用或冲销", "性能并发可用性SLA"
    ]),
}


def section_appendices(doc):
    add_heading(doc, "附录A：功能需求编号矩阵", 1)
    add_para(doc, "优先级建议：P0为上线必需闭环，P1为阶段内重要能力，P2为优化或高级能力。以下矩阵可作为供应商逐项响应表的基础，最终应补充产品版本、标准/配置/定制、交付阶段、责任方和合同条款。")
    rows = []
    for prefix, (domain, items) in REQ_GROUPS.items():
        for idx, req in enumerate(items, start=1):
            priority = "P0" if idx <= max(3, int(len(items) * 0.55)) else ("P1" if idx <= int(len(items) * 0.9) else "P2")
            owner = {
                "GOV": "项目管理", "ORG": "平台/组织权限", "OTL": "ERP/工作流/MES", "MDM": "ERP/PLM", "QTN": "ERP/CPQ", "HKO": "ERP/外贸/财务", "PMC": "ERP/APS", "SRM": "SRM/ERP", "MKT": "行情数据/ERP", "MES": "MES", "OPC": "ERP成本/MES",
                "QMS": "QMS", "WMS": "WMS", "FIN": "ERP财务", "ADM": "行政/考勤", "DOC": "报表中心", "ALT": "工作流/预警", "CUS": "客户门户",
                "BI": "BI", "AIX": "AI分析平台", "SEC": "平台/安全",
            }[prefix]
            acceptance = "真实业务演示并保留单据、版本、日志或报表证据"
            rows.append((f"{prefix}-{idx:02d}", domain, priority, owner, req, acceptance))
    add_table(doc, ["编号", "领域", "优先级", "承载系统", "详细需求", "验收证据"], rows, [900, 1000, 700, 1200, 3400, 2160], font_size=7.2)

    add_heading(doc, "附录B：系统生成文件清单", 1)
    rows2 = [(str(i + 1), name, trigger, output) for i, (name, trigger, _data, output) in enumerate(REPORT_SPECS)]
    add_table(doc, ["序号", "文件/报表", "生成时机", "输出及控制"], rows2, [650, 1800, 3100, 3810], font_size=8.2)

    add_heading(doc, "附录C：选型演示问题清单", 1)
    questions = [
        "请分别建立收费模具订单、免费样品订单和正式业务订单，证明模具/样品可按审批免费但仍归集成本，正式业务订单零价或缺付款条件时被系统阻断，并展示样品转正式订单及费用抵扣履历。",
        "请在客户主数据勾选一个香港代生产价格客户，建立正式、样品和模具订单并录入同一原始价格；证明只有正式订单自动按70%生成订单价格，另外两类保持原价。",
        "请将同一订单切换到未勾选客户、修改订单类型和原始价格，展示自动重判、重算、版本、重新审核和防止重复乘70%；随后按普通流程完成MRP、生产、报关、开票和收款。",
        "请使用我方真实图纸展示相似产品、历史报价、采购/委外价格和成本分析，并生成客户报价PDF。",
        "请展示交易所/持牌金属行情、供应商报价和历史采购价如何映射、换算为企业落地价，并分别向采购、财务和业务展示授权视图。",
        "请展示材料价格快照如何进入报价和标准成本模拟，行情超阈值如何预警且不会自动篡改已审批订单或法定存货账。",
        "请展示财务、总经办、业务、工程、PMC、采购、委外、仓库、生产、品质、后工序、行政和IT十三部门的独立工作台、数据范围、待办和跨部门交接。",
        "请展示CNC/车床工件如何逐站进入清洗、丝印、镭雕、打磨、振磨和全检包装，并将组装作为独立BOM、齐套、工单、测试和成本管理。",
        "请用一张真实订单展示切料、CNC/车床、打磨、清洗、委外表处、镭雕、丝印、组装、全检和包装分别建立工序成本卡，本工序新增成本如何逐站累计为订单实际成本。",
        "请展示工序成本的标准/实际材料、人工、机时、刀具工装、委外、检验、制造费用和差异；证明返工重做会生成新工序实例并再次计费，而不是覆盖或沿用原成本。",
        "请展示首检、巡检/点检、末件检和FAI如何自动触发、保存原始测量值与版本并生成系统报告。",
        "请展示考勤机/门禁数据、班次、请假加班、异常、MES工时差异和两个不同劳务单位的对账结算。",
        "请展示图纸变更如何联动BOM、工艺、检验、库存、采购、在制和委外，不接受只上传附件。",
        "请展示MRP净需求的完整计算依据以及现有库存被其他订单占用时的齐套变化。",
        "请展示多个订单共用设备、夹具、技术员或检具时的冲突、排程和交期模拟。",
        "请展示第五工序退回第二工序返工、复检后直接返回第五工序的路线，说明第三、第四工序原结果如何确认有效。",
        "请展示第五工序退回第二工序返工后，必须按第三、第四、第五工序顺序重流的路线，证明系统不能直接跳回第五工序且原履历不被删除。",
        "请展示同一异常批次按数量拆分为直接返回和顺序重流两条路线，子批次、二维码、在制、检验、成本和最终汇合如何管理。",
        "请演示包装发现异常且责任为CNC：系统拆出返工子订单，CNC返工后再次经过打磨、委外表处、镭雕、丝印、全检和包装；证明每道重复工序重新计费、第二次表处正常形成供应商应付，并将全部返工毛成本归集到CNC责任。",
        "请演示包装发现异常且责任为委外表处供应商：系统生成无偿返工单并将供应商追加应付控制为零，同时完整核算退运、复检、后续重做、包装、停线和延期等供应商造成的毛损失、索赔回收及企业净损失。",
        "请展示不允许跳工序和有条件跳工序的差异及审批日志。",
        "请展示供应商在网页报价、议价、定标、确认交期、创建ASN、打印送货单并完成对账。",
        "请展示客户门户的项目里程碑、交货计划和文件，同时证明客户无法看到内部成本和其他客户。",
        "请展示工序、审批、采购、委外、质量和设备超期预警的升级和关闭。",
        "请展示报价、订单/工序/采购/齐套/供应商进度、逐工序成本卡/成本瀑布/标准实际差异/在制成本、香港客户正式订单70%价格审计、报关、对账、品质、停机、报废、良率和呆滞报表由系统生成。",
        "请使用固定财务问题展示AI对应收、资金、成本差异、毛利和月结异常的问答，答案必须显示口径、截止时间并引用凭证或业务单据。",
        "请展示老板通过AI追问经营指标、交付与现金风险并进行情景模拟，同时证明AI不能越权或直接执行记账、付款、关账和审批。",
        "请说明每项能力是标准、配置、增购、定制还是第三方接口，并提供版本、工作量和升级影响。",
        "请提供接口失败、网络中断、重复消息、数据补传和对账的处理方案。",
    ]
    add_numbered(doc, questions, font_size=9, space_after=2, line_spacing=1.1)

    add_heading(doc, "附录D：厂商、行情及关务财税官方资料参考", 1)
    refs = [
        ("鼎捷ERP产品线", "https://www.digiwin.com/software/digiwinERP.html"),
        ("鼎捷APS", "https://www.digiwin.com/project/APS/APS.html"),
        ("鼎捷sQMS", "https://www.digiwin.com/project/sQMS/sQMS.html"),
        ("用友U9 cloud", "https://u9cloud.yonyou.com/"),
        ("用友YonSuite", "https://www.yonyou.com/global/products/yonsuite/"),
        ("金蝶云·星空", "https://www.kingdee.com/products/galaxy.html"),
        ("金蝶制造云", "https://www.kingdee.com/products/galaxy_manufacture.html"),
        ("金蝶智慧工厂云", "https://www.kingdee.com/products/galaxy_smart_factory.html"),
        ("黑湖小工单", "https://www.xiaogongdan.cn/products"),
        ("上海期货交易所官网", "https://www.shfe.com.cn/"),
        ("伦敦金属交易所市场数据", "https://www.lme.com/market-data"),
        ("伦敦金属交易所数据许可", "https://www.lme.com/Market-data/Market-data-licensing"),
        ("海关进出口货物申报办事指南", "https://online.customs.gov.cn/static/pages/guides/000629002001/000629002001.html"),
        ("国家税务总局公告2016年第42号：关联申报和同期资料", "https://fgk.chinatax.gov.cn/zcfgk/c100012/c5194659/content.html"),
        ("国家外汇管理局汇发〔2024〕11号：优化贸易外汇业务管理", "https://m.safe.gov.cn/safe/2024/0407/24204.html"),
    ]
    add_table(doc, ["资料", "链接"], refs, [2300, 7060], font_size=8.2)
    add_para(doc, "注：功能与许可会调整，以项目时点的书面响应、合同、演示和验收为准。", italic=True, after=0)


def build():
    doc = setup_document()
    add_cover(doc)
    add_front_matter(doc)
    section_overview(doc)
    section_architecture(doc)
    section_order_lifecycle(doc)
    section_department_blueprint(doc)
    section_roadmap(doc)
    section_master_engineering(doc)
    section_sales_customer(doc)
    section_procurement_supplier(doc)
    section_planning(doc)
    section_mes(doc)
    section_qms(doc)
    section_wms(doc)
    section_finance(doc)
    section_documents(doc)
    section_alerts(doc)
    section_portal_security(doc)
    section_bi(doc)
    section_integration(doc)
    section_vendor(doc)
    section_implementation(doc)
    section_acceptance(doc)
    section_appendices(doc)

    # Clean trailing paragraphs and ensure stable widow/orphan behavior.
    for p in doc.paragraphs:
        p_pr = p._p.get_or_add_pPr()
        if p.style and p.style.name.startswith("Heading"):
            keep_next = p_pr.find(qn("w:keepNext"))
            if keep_next is None:
                keep_next = OxmlElement("w:keepNext")
                p_pr.append(keep_next)
        widow = p_pr.find(qn("w:widowControl"))
        if widow is None:
            widow = OxmlElement("w:widowControl")
            widow.set(qn("w:val"), "true")
            p_pr.append(widow)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
